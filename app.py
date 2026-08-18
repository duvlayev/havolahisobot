from flask import Flask, request, jsonify, send_from_directory, session
import sqlite3
import os
import re
import math
import uuid
import json
import threading
import time as time_module
from functools import wraps
from datetime import datetime, timedelta
import openpyxl
from io import BytesIO
from flask import send_file
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from dotenv import load_dotenv
import requests
from urllib.parse import urlparse

# --- .env fayldan sozlamalarni o'qish ---
# MUHIM: Telegram bot TOKEN va guruh ID kodga yozilmaydi — ular serverdagi
# ".env" faylida (app.py bilan bir papkada) saqlanadi. Namuna uchun
# ".env.example" faylga qarang. Fayl topilmasa — bot funksiyalari o'chirilgan
# holda ishlaydi (dastur qulamaydi, faqat papkaoch bo'limida ogohlantirish
# ko'rsatiladi).
load_dotenv()
TELEGRAM_BOT_TOKEN = os.environ.get('TELEGRAM_BOT_TOKEN', '').strip()
TELEGRAM_GROUP_ID = os.environ.get('TELEGRAM_GROUP_ID', '').strip()
# Ixtiyoriy: xatoliklar shu shaxsiy chatga yuboriladi (bo'lmasa — faqat logga
# va veb-sahifadagi "Oxirgi faoliyat" ro'yxatiga yoziladi).
TELEGRAM_ADMIN_CHAT_ID = os.environ.get('TELEGRAM_ADMIN_CHAT_ID', '').strip()
TELEGRAM_API_BASE = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}" if TELEGRAM_BOT_TOKEN else None

def papkaoch_bot_configured():
    return bool(TELEGRAM_BOT_TOKEN and TELEGRAM_GROUP_ID)

app = Flask(__name__, static_folder='static')
DB = 'links.db'

# --- Admin autentifikatsiyasi ---
# DIQQAT: production serverda quyidagi ikkala qiymatni ham albatta o'zgartiring
# (masalan environment o'zgaruvchilar orqali: SECRET_KEY, ADMIN_PASSWORD)
app.secret_key = os.environ.get('SECRET_KEY', 'buni-production-da-ozgartiring-2026')
ADMIN_PASSWORD = os.environ.get('ADMIN_PASSWORD', 'S07h09m14')
app.permanent_session_lifetime = timedelta(hours=12)

# --- Screenshot fayllari uchun sozlamalar ---
UPLOAD_SUBDIR = os.path.join('uploads', 'screenshots')  # static ichida
UPLOAD_DIR = os.path.join(app.static_folder, UPLOAD_SUBDIR)
os.makedirs(UPLOAD_DIR, exist_ok=True)
ALLOWED_SHOT_EXT = {'png', 'jpg', 'jpeg', 'webp', 'gif'}

def allowed_shot_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_SHOT_EXT

SHOT_COLUMNS = {'before': 'screenshot_before', 'after': 'screenshot_after'}

def admin_required(f):
    """Faqat admin sifatida tizimga kirgan foydalanuvchilarga ruxsat beruvchi dekorator."""
    @wraps(f)
    def wrapper(*args, **kwargs):
        if not session.get('is_admin'):
            return jsonify({'error': "Ruxsat yo'q. Bu amalni faqat admin bajara oladi."}), 403
        return f(*args, **kwargs)
    return wrapper

# Bitta bazada bir nechta "loyiha" (board) saqlanadi — har biri o'z sahifasidek
# ishlaydi, lekin bir xil URL/server orqali, faqat "board" ustuni bilan ajratiladi.
BOARDS = {
    'jaloba':    "Jaloba",
    'fishing':   "Fishing",
    't_hisobot': "T.hisobot",
}
DEFAULT_BOARD = 'jaloba'

# T.hisobot boshqa loyihalardan farqli status ro'yxatiga ega (Faol/O'chgan o'rniga
# rukunlar). Har bir board uchun ruxsat etilgan status qiymatlari va standart qiymati.
STATUS_OPTIONS = {
    'jaloba':    ['faol', "o'chgan"],
    'fishing':   ['faol', "o'chgan"],
    't_hisobot': ['Umumiy', 'Tashrif', 'Videoselektor', 'Taqdimot', 'Uchrashuv', 'Tabrik', 'Saylov'],
}
DEFAULT_STATUS = {
    'jaloba':    'faol',
    'fishing':   'faol',
    't_hisobot': 'Umumiy',
}

def _detect_sana_from_path(folder_path, max_levels=6):
    """Berilgan papka yo'lidagi ota-papkalarni (barglardan yuqoriga qarab, ENG
    YAQINIDAN boshlab) "DD.MM.YYYY" nomli papka uchun tekshiradi va birinchi mos
    kelganini "YYYY-MM-DD" formatida qaytaradi. Hech biri mos kelmasa — None.

    MUHIM: faqat bevosita ota-papkani emas — bir necha daraja yuqorigacha
    (standart: 6 daraja) tekshiramiz, chunki sana papkasi bilan maqsad papka
    ("Жалоба", "Фишинг", "Коммент" va h.k.) orasida qo'shimcha oraliq
    papka(lar) bo'lishi mumkin, masalan:
        ".../05.08.2026/Жалоба"                    -> "05.08.2026" (1-daraja)
        ".../05.08.2026/Biror papka nomi/Жалоба"    -> "05.08.2026" (2-daraja)
    """
    current = os.path.normpath(folder_path)
    for _ in range(max_levels):
        parent = os.path.dirname(current)
        if not parent or parent == current:
            break
        name = os.path.basename(parent)
        try:
            return datetime.strptime(name, '%d.%m.%Y').strftime('%Y-%m-%d')
        except ValueError:
            current = parent
    return None

def valid_status_for_board(board, status):
    """Berilgan status shu board uchun ruxsat etilganmi — bo'lmasa standart qiymatga tushadi.

    MUHIM: lotin va kirill alifbosida yozilgan bir xil ma'noli statuslar
    (masalan "Tashrif" va "Ташриф") bir xil status sifatida tan olinadi —
    CATEGORY_FOLDER_TO_STATUS lug'ati orqali (pastda e'lon qilingan).

    T.HISOBOT UCHUN MUHIM: status ro'yxati ENDI OCHIQ — "Коммент" papkasi
    ichida istalgan nomdagi (oldindan belgilanmagan) rukun papkalari ham
    bo'lishi mumkin (masalan "Таълим", yoki "Sanoat / ilmiymaqola" kabi
    bir necha daraja ichma-ich), shuning uchun bunday nomlar ham qabul
    qilinadi — faqat 6 ta "taniqli" nom (Umumiy/Tashrif/...) kanonik
    shaklga keltiriladi, qolganlari o'z original nomicha qoladi."""
    options = STATUS_OPTIONS.get(board, STATUS_OPTIONS[DEFAULT_BOARD])
    s = (status or '').strip()
    if s:
        # 1) To'g'ridan-to'g'ri mos kelishi (katta-kichik harfga sezgir emas)
        for opt in options:
            if s.lower() == opt.lower():
                return opt
        # 2) T.hisobot uchun — lotin/kirill alifbosidagi variantlarni bir xil
        #    status sifatida taniymiz (masalan "Ташриф" -> "Tashrif")
        if board == 't_hisobot':
            mapped = CATEGORY_FOLDER_TO_STATUS.get(s.lower())
            if mapped:
                return mapped
            # 2b) Boshqa (oldindan belgilanmagan, papka tuzilishidan kelib
            # chiqqan) status nomlari ham — original holicha — qabul qilinadi.
            return s
        # 3) Jaloba/Fishing uchun eski/muqobil yozilishlar bilan moslik
        else:
            if s.lower() in ("o'chgan", 'ochgan', "o`chgan", 'o‘chgan', 'inactive', 'off', '0', 'false'):
                return "o'chgan"
            if s.lower() in ('faol', 'active', 'on', '1', 'true'):
                return 'faol'
    return DEFAULT_STATUS.get(board, options[0])

# "Jalobalar soni (min/max)" filtri qaysi ustunga nisbatan ishlashi — T.hisobot uchun
# "Umumiy izohlar soni" ustuni ishlatiladi.
COUNT_COLUMN = {
    'jaloba':    'likes',
    'fishing':   'likes',
    't_hisobot': 'umumiy_izohlar',
}

def get_db():
    conn = sqlite3.connect(DB)
    conn.row_factory = sqlite3.Row
    return conn

def get_setting(key, default=None):
    with get_db() as conn:
        row = conn.execute('SELECT value FROM settings WHERE key=?', (key,)).fetchone()
    return row['value'] if row is not None else default

def set_setting(key, value):
    with get_db() as conn:
        conn.execute(
            'INSERT INTO settings (key, value) VALUES (?, ?) '
            'ON CONFLICT(key) DO UPDATE SET value=excluded.value',
            (key, value)
        )
        conn.commit()

# ============================================================================
# TARMOQ PAPKASIDAGI SCREENSHOTLARNI HISOBLASH
# ("main_hisoblash_.py" dasturidagi mantiq bilan bir xil algoritm)
#
# Kutilgan papka tuzilishi (tarmoq/NAS diskda):
#   <root_path>/29.07.2026/Жалоба/1. t.me (tm)/screenshot1.png
#   <root_path>/29.07.2026/Жалоба/1. t.me (tm)/Коммент/screenshot2.png
#   <root_path>/29.07.2026/Жалоба/2. tiktok.com (tk)/...
#   <root_path>/29.07.2026/Фишинг/1. .../...
#
# Har bir "N. domen (kod)" papkasi bitta havolaga mos keladi (N — o'sha
# kungi Telegram postidagi tartib raqami). "Коммент" papkasi ichidagi
# rasmlar asosiy sondan alohida ("comment_count") hisoblanadi.
# ============================================================================

SCREENSHOT_EXTENSIONS = ('.jpg', '.jpeg', '.png', '.webp', '.bmp', '.gif')

# Board -> tarmoqdagi maxsus papka nomi (bir nechta variant, lotin/kirill)
BOARD_SPECIAL_FOLDER_NAMES = {
    'jaloba':  ('Жалоба', 'Jaloba'),
    'fishing': ('Фишинг', 'Fishing'),
}

# T.hisobot uchun: papka manzili "Коммент" (yoki variantlari) bilan tugaydi
COMMENT_FOLDER_NAMES = ('Коммент', 'коммент', 'Комент', 'комент', 'Komment', 'komment')

# "Коммент" papkasi ichidagi rukun-nomli pastki papkalar (kirill/lotin) ->
# bizning ichki status qiymatimiz. Kichik harfga o'tkazib solishtiriladi.
CATEGORY_FOLDER_TO_STATUS = {
    'умумий': 'Umumiy', 'umumiy': 'Umumiy',
    'ташриф': 'Tashrif', 'tashrif': 'Tashrif',
    'видеоселектор': 'Videoselektor', 'videoselektor': 'Videoselektor',
    'тақдимот': 'Taqdimot', 'такдимот': 'Taqdimot', 'taqdimot': 'Taqdimot',
    'учрашув': 'Uchrashuv', 'uchrashuv': 'Uchrashuv',
    'табрик': 'Tabrik', 'tabrik': 'Tabrik',
    'сайлов': 'Saylov', 'saylov': 'Saylov',
}

def _discover_t_hisobot_category_dirs(root_path, path_prefix=None, max_depth=8):
    """"Коммент" papkasi ichida — istalgan chuqurlikdagi (ota/bola/nabira...)
    pastki papkalarda — raqamlangan post-papkalarga (masalan "1. Facebook
    (fb)") EGA bo'lgan har bir "barg" (leaf) papkani alohida rukun/bo'lim
    sifatida topadi. Ichma-ich papkalar bir nechta va turli nomda bo'lishi
    mumkin — ular "Ota / Bola / Nabira" ko'rinishida BITTA birlashtirilgan
    rukun nomiga aylantiriladi.

    Oldindan belgilangan 6 ta rukun nomi (Умумий/Ташриф/Видеоселектор/...)
    — lotin/kirill, katta/kichik harf farqisiz — hamon o'zining KANONIK
    nomiga (Umumiy/Tashrif/...) keltiriladi, bu mavjud bazadagi eski
    yozuvlar bilan mosligini saqlaydi. Boshqa (yangi, ixtiyoriy) nomlar esa
    original ko'rinishicha qoladi.

    Qaytaradi: [(rukun_nomi, to'liq_yo'l), ...] — papka nomlari bo'yicha
    alifbo tartibida, topilgan barcha "barg" papkalar uchun."""
    if path_prefix is None:
        path_prefix = []
    if len(path_prefix) >= max_depth:
        return []
    try:
        entries = sorted(os.listdir(root_path))
    except OSError:
        return []

    subdirs = [(name, os.path.join(root_path, name)) for name in entries
               if os.path.isdir(os.path.join(root_path, name))]

    # Shu darajaning o'zida raqamlangan post-papkalar (masalan "1. t.me (tm)")
    # bormi? Bo'lsa — bu "barg" (leaf) papka, ya'ni alohida rukun.
    has_numbered = any(_extract_link_folder_order(name) is not None for name, _ in subdirs)
    if has_numbered:
        label = ' / '.join(path_prefix) if path_prefix else 'Umumiy'
        return [(label, root_path)]

    if not subdirs:
        return []

    # Aks holda — bu oraliq guruh papkasi, pastkilarga tushamiz
    results = []
    for name, full in subdirs:
        canonical = CATEGORY_FOLDER_TO_STATUS.get(name.strip().lower())
        next_prefix = path_prefix + [canonical or name]
        results.extend(_discover_t_hisobot_category_dirs(full, next_prefix, max_depth))
    return results

def _screenshot_plus_count(filename):
    """Fayl nomidan '+N' yoki 'N+' naqshini o'qib, bitta fayl nechta
    jalobani anglatishini aniqlaydi (topilmasa — 1ta).

    Bu funksiya "main(hisoblash)_new.py" dagi extract_info funksiyasidagi
    hisoblash mantig'i bilan AYNAN bir xil ishlaydi: fayl nomi ichida
    "raqam+" yoki "+raqam" ko'rinishidagi barcha birikmalar topiladi (bo'sh
    joy bilan ajratilgan yoki ismga yopishgan bo'lishidan qat'i nazar) va
    ularning qiymatlari qo'shib chiqiladi. "-" belgisi hech qachon hisobga
    olinmaydi. Hech qanday "+" topilmasa — natija 1ta.
    """
    name_part = os.path.splitext(filename)[0]
    plus_matches = re.findall(r'(\d+)\s*\+|\+\s*(\d+)', name_part)
    count = 0
    for left, right in plus_matches:
        if left and left.isdigit():
            count += int(left)
        elif right and right.isdigit():
            count += int(right)
    return count if count > 0 else 1

def _extract_link_folder_order(folder_name):
    """Papka nomidan boshidagi tartib raqamini ajratib oladi:
    '1. t.me (tm)' -> 1, '12' -> 12, 'Коммент' -> None."""
    folder_name = folder_name.strip()
    if folder_name.isdigit():
        return int(folder_name)
    m = re.match(r'^(\d+)[.\s]', folder_name)
    if m:
        return int(m.group(1))
    return None

def _count_link_folder_screenshots(folder_path):
    """Bitta 'N. domen (kod)' papkasi ichidagi rasmlarni sanaydi.
    Qaytaradi: (main_count, comment_count).
    'Коммент' nomli PASTKI papka (folder_path ICHIDA) alohida hisoblanadi.
    Muhim: folder_path ning o'zi biror 'Коммент' OTA-papka ichida joylashgan
    bo'lishi mumkin (T.hisobot holatida shunday) — bu ota-papka nomlari
    hisobga olinmaydi, faqat folder_path'dan PASTdagi yo'l qismlari tekshiriladi."""
    main_count = 0
    comment_count = 0
    base = os.path.normpath(folder_path)
    for root, _, files in os.walk(base):
        rel = os.path.relpath(root, base)
        rel_parts = [] if rel == '.' else rel.split(os.sep)
        under_comment = any(p.lower() in ('коммент', 'komment') for p in rel_parts)
        for f in files:
            if f.lower().endswith(SCREENSHOT_EXTENSIONS):
                c = _screenshot_plus_count(f)
                if under_comment:
                    comment_count += c
                else:
                    main_count += c
    return main_count, comment_count

def _find_board_special_folder(date_folder_path, board):
    """Sana papkasi ichidan 'Жалоба'/'Фишинг' (yoki lotin variantini) qidiradi."""
    for name in BOARD_SPECIAL_FOLDER_NAMES.get(board, ()):
        p = os.path.join(date_folder_path, name)
        if os.path.isdir(p):
            return p
    return None

def guess_folder_category(folder_name):
    """Papka nomidagi belgi ((fb)/(ins)/(yt)/(tm)/(tw)/(tk) va kirill variantlari)
    bo'yicha ijtimoiy tarmoq turini taxmin qiladi — faqat hodimga vizual
    tekshirish uchun ko'rsatiladi, moslashtirishning o'zi buning ustiga qurilmagan."""
    if not folder_name:
        return None
    if re.search(r'\([fF][bB]\)|\([фФ][бБ]\)', folder_name):
        return 'Facebook'
    if re.search(r'\([iI][nN][sS]\)|\([иИ][нН][сС]\)', folder_name):
        return 'Instagram'
    if re.search(r'\([yY][tT]\)|\([юЮ][тТ]\)', folder_name):
        return 'YouTube'
    if re.search(r'\([tT][mM]\)|\([тТ][мМ]\)', folder_name):
        return 'Telegram'
    if re.search(r'\([tT][wW]\)|\([тТ][вВ]\)', folder_name):
        return 'Twitter'
    if re.search(r'\([tT][kK]\)|\([тТ][кК]\)', folder_name):
        return 'TikTok'
    return 'Boshqalar'

def guess_url_category(url):
    """Havola domeni bo'yicha ijtimoiy tarmoq turini taxmin qiladi."""
    if not url:
        return None
    u = url.lower()
    if 't.me' in u or 'telegram.me' in u or 'telegram.org' in u:
        return 'Telegram'
    if 'tiktok.com' in u:
        return 'TikTok'
    if 'youtube.com' in u or 'youtu.be' in u:
        return 'YouTube'
    if 'facebook.com' in u or 'fb.com' in u or 'fb.watch' in u:
        return 'Facebook'
    if 'instagram.com' in u:
        return 'Instagram'
    if 'twitter.com' in u or 'x.com' in u:
        return 'Twitter'
    return 'Boshqalar'

# ESLATMA: bu yerda avval "fon rejimida avtomatik sinxronlash" (har necha daqiqada
# o'zi tekshirib, ma'lumotlarni jimgina yangilab turuvchi) mexanizmi bo'lgan edi.
# U olib tashlandi, chunki talab o'zgardi: endi ma'lumotlar bazaga FAQAT hodim
# "/api/screenshot-folder/preview" orqali papkani ko'rib chiqib, xatoliklarni
# tuzatib, keyin "/api/screenshot-folder/apply" orqali TASDIQLAGANIDAN so'ng
# yoziladi — va shundan keyin hodim qayta shu jarayonni bajarmaguncha o'zgarmay
# turadi. Bu inson tekshiruvisiz avtomatik yozishning oldini oladi.

def init_db():
    with get_db() as conn:
        conn.execute('''
            CREATE TABLE IF NOT EXISTS links (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                url TEXT NOT NULL,
                sana TEXT NOT NULL,
                likes INTEGER DEFAULT 0,
                status TEXT DEFAULT "faol",
                board TEXT NOT NULL DEFAULT "jaloba",
                created_at TEXT DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        # Eski bazada "board" ustuni bo'lmasa (avvalgi versiyadan yangilanganda),
        # xatolikka uchramasdan qo'shib qo'yamiz.
        try:
            conn.execute('ALTER TABLE links ADD COLUMN board TEXT NOT NULL DEFAULT "jaloba"')
        except sqlite3.OperationalError:
            pass
        # Screenshot ustunlari: dastlabki (faol paytdagi) va keyingi (o'chgan paytdagi)
        try:
            conn.execute('ALTER TABLE links ADD COLUMN screenshot_before TEXT')
        except sqlite3.OperationalError:
            pass
        try:
            conn.execute('ALTER TABLE links ADD COLUMN screenshot_after TEXT')
        except sqlite3.OperationalError:
            pass
        # Ma'lumot o'zgartirilgan sana — faqat qatordagi biror maydon tahrirlanganda to'ldiriladi
        try:
            conn.execute('ALTER TABLE links ADD COLUMN edited_at TEXT')
        except sqlite3.OperationalError:
            pass
        # T.hisobot bo'limi uchun: umumiy va ajratilgan izohlar soni
        try:
            conn.execute('ALTER TABLE links ADD COLUMN umumiy_izohlar INTEGER DEFAULT 0')
        except sqlite3.OperationalError:
            pass
        try:
            conn.execute('ALTER TABLE links ADD COLUMN ajratilgan_izohlar INTEGER DEFAULT 0')
        except sqlite3.OperationalError:
            pass
        # Havolaning Telegram postidagi tartib raqami (masalan "#Жалоба" ro'yxatidagi
        # 1, 2, 3...). Har kuni 1 dan boshlanadi (board+sana bo'yicha alohida hisoblanadi).
        # Tarmoq papkasidagi "1. t.me (tm)" kabi papka nomlariga mos kelishi uchun kerak.
        try:
            conn.execute('ALTER TABLE links ADD COLUMN order_num INTEGER')
        except sqlite3.OperationalError:
            pass
        # Bir xil link (URL) keyingi kunlarda qayta skanerlansa, oldingi kunlargi
        # sonini YO'QOTMASLIK uchun: "joriy" son (likes / umumiy_izohlar /
        # ajratilgan_izohlar) har doim ENG SO'NGGI kunning sonini bildiradi, bu
        # ustunlar esa o'sha linkka OLDINGI kunlarda (bugungidan oldin) yig'ilgan
        # jami sonni saqlaydi. Masalan: 1-kun 50, 2-kun 10 bo'lsa -> likes=10,
        # prev_likes=50 (ekranda "(50) 10" ko'rinishida chiqadi).
        try:
            conn.execute('ALTER TABLE links ADD COLUMN prev_likes INTEGER DEFAULT 0')
        except sqlite3.OperationalError:
            pass
        try:
            conn.execute('ALTER TABLE links ADD COLUMN prev_umumiy_izohlar INTEGER DEFAULT 0')
        except sqlite3.OperationalError:
            pass
        try:
            conn.execute('ALTER TABLE links ADD COLUMN prev_ajratilgan_izohlar INTEGER DEFAULT 0')
        except sqlite3.OperationalError:
            pass
        # Sozlamalar (masalan tarmoq papkasi manzili, oxirgi sinxronlash vaqti) uchun
        # oddiy key-value jadval.
        conn.execute('''
            CREATE TABLE IF NOT EXISTS settings (
                key TEXT PRIMARY KEY,
                value TEXT
            )
        ''')

        # --- "Papkaoch" bo'limi uchun jadvallar ---

        # Domenlar bazasi: Telegram guruhga tashlangan havolalarning asosiy
        # domenini aniqlash uchun. Admin bu ro'yxatni to'ldirishi, o'zgartirishi
        # va o'chirishi mumkin (veb-sahifada).
        conn.execute('''
            CREATE TABLE IF NOT EXISTS papkaoch_domains (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                base_url TEXT NOT NULL,
                short_code TEXT NOT NULL,
                created_at TEXT
            )
        ''')
        # Standart domenlar — faqat jadval BO'SH bo'lsa bir marta to'ldiriladi
        # (talab hujjatidagi 4-bo'limdagi jadvaldan).
        existing_count = conn.execute('SELECT COUNT(*) FROM papkaoch_domains').fetchone()[0]
        if existing_count == 0:
            default_domains = [
                ('Facebook', 'https://www.facebook.com/', 'fb'),
                ('Instagram', 'https://www.instagram.com/', 'ins'),
                ('Youtube', 'https://www.youtube.com/', 'yt'),
                ('Tiktok', 'https://www.tiktok.com/', 'tk'),
                ('Telegram', 'https://www.t.me/', 'tm'),
                ('VKontakte', 'https://vk.ru/', 'vk'),
                ('X', 'https://x.com/', 'tw'),
                ('LinkedIn', 'https://linkedin.com/', 'ln'),
            ]
            now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            for name, base_url, short_code in default_domains:
                conn.execute(
                    'INSERT INTO papkaoch_domains (name, base_url, short_code, created_at) VALUES (?,?,?,?)',
                    (name, base_url, short_code, now)
                )

        # Har bir qayta ishlangan (yaratilgan/rad etilgan/xato/kutilayotgan)
        # havola-papka jufti shu yerda qayd etiladi — takroriylikni tekshirish
        # va veb-sahifadagi "Oxirgi faoliyat" ro'yxati uchun ishlatiladi.
        conn.execute('''
            CREATE TABLE IF NOT EXISTS papkaoch_entries (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                chat_id TEXT,
                message_id INTEGER,
                tags_json TEXT NOT NULL,
                order_num INTEGER,
                url TEXT NOT NULL,
                domain_name TEXT,
                short_code TEXT,
                folder_name TEXT,
                full_path TEXT,
                status TEXT NOT NULL,
                note TEXT,
                created_at TEXT
            )
        ''')
        conn.execute('CREATE INDEX IF NOT EXISTS idx_papkaoch_entries_url ON papkaoch_entries(url)')

        # "Boshqa ID/papka bilan takror tashlansa" holatida Telegram guruhda
        # HA/YO'Q tugmalari bilan tasdiq so'ralganda, tugma bosilguncha kutilib
        # turadigan holat shu yerda saqlanadi (callback_data orqali topiladi).
        conn.execute('''
            CREATE TABLE IF NOT EXISTS papkaoch_pending (
                token TEXT PRIMARY KEY,
                chat_id TEXT,
                prompt_message_id INTEGER,
                tags_json TEXT NOT NULL,
                order_num INTEGER,
                url TEXT NOT NULL,
                domain_name TEXT,
                short_code TEXT,
                folder_name TEXT,
                full_path TEXT,
                created_at TEXT,
                resolved INTEGER DEFAULT 0
            )
        ''')
        conn.commit()

    # Bir martalik migratsiya: eski qatorlarda lotin/kirill aralash yozilgan status
    # qiymatlarini (masalan "Ташриф") bir xil kanonik qiymatga ("Tashrif")
    # keltiramiz — aks holda faqat BUNDAN KEYIN qo'shiladigan yozuvlar to'g'ri
    # normalizatsiya qilinib, eski yozuvlar noto'g'ri (yoki alohida) status
    # sifatida qolib ketardi. Har safar ilova ishga tushganda ishlaydi, lekin
    # allaqachon kanonik qiymatga ega qatorlarga tegmaydi (tez va xavfsiz).
    with get_db() as conn:
        rows = conn.execute("SELECT id, board, status FROM links").fetchall()
        for r in rows:
            normalized = valid_status_for_board(r['board'], r['status'])
            if normalized != r['status']:
                conn.execute('UPDATE links SET status=? WHERE id=?', (normalized, r['id']))
        conn.commit()

def get_board(args):
    """So'rovdan board qiymatini oladi, noto'g'ri/berilmagan bo'lsa standart qiymat qaytadi."""
    board = args.get('board', DEFAULT_BOARD)
    return board if board in BOARDS else DEFAULT_BOARD

# ============================================================================
# "PAPKAOCH" BO'LIMI — Telegram guruhga tashlangan havolalar bo'yicha
# avtomatik papka yaratish
# ============================================================================
#
# Bu bo'lim quyidagi vazifalarni bajaradi:
#   1) Telegram xabarini tahlil qilib, teglar (#ota_papka #bola_papka ...) va
#      raqamlangan havolalar ro'yxatini ajratib oladi.
#   2) Har bir havolaning domenini bazadagi (papkaoch_domains) ro'yxat bilan
#      solishtiradi va papka nomini shakllantiradi.
#   3) Takroriylikni tekshiradi (xuddi shu joyda/raqamda — rad etiladi;
#      boshqa joyda/raqamda — Telegram orqali HA/YO'Q so'raladi).
#   4) Papkani serverdagi belgilangan (root) manzilda yaratadi.
#   5) Natijani guruhga (✅/❌) va log faylga yozadi.
#
# QUYIDAGI FUNKSIYALARNING KO'PCHILIGI TARMOQQA BOG'LIQ EMAS (faqat matn
# tahlili, fayl tizimi va SQLite) — shuning uchun ular sinov (test) muhitida
# to'liq tekshirilgan. Faqat Telegram API bilan bevosita ishlaydigan
# funksiyalar (tg_*, papkaoch_bot_loop) real bot TOKEN/GURUH ID talab qiladi
# va faqat haqiqiy serverda sinaladi.

PAPKAOCH_LOG_DIR = 'papkaoch_logs'
os.makedirs(PAPKAOCH_LOG_DIR, exist_ok=True)

def papkaoch_log(tag_path, url, result, extra=''):
    """Har bir jarayonni kunlik log faylga yozadi.
    Format: havola_log_(DD-MM.YYYY).txt
    Yozuv: [vaqt] | [xeshteg] | [havola] | [natija]"""
    now = datetime.now()
    fname = now.strftime('havola_log_(%d-%m.%Y).txt')
    fpath = os.path.join(PAPKAOCH_LOG_DIR, fname)
    tag_str = ' -> '.join(f'#{t}' for t in tag_path) if tag_path else '—'
    line = f"[{now.strftime('%Y-%m-%d %H:%M:%S')}] | {tag_str} | {url} | {result}"
    if extra:
        line += f" | {extra}"
    try:
        with open(fpath, 'a', encoding='utf-8') as f:
            f.write(line + '\n')
    except OSError:
        pass  # log yozib bo'lmasa ham asosiy jarayon to'xtamasin

def papkaoch_list_log_dates():
    """Mavjud log fayllarning sanalarini ('YYYY-MM-DD' formatida) ro'yxat qilib qaytaradi."""
    dates = []
    if not os.path.isdir(PAPKAOCH_LOG_DIR):
        return dates
    for fname in os.listdir(PAPKAOCH_LOG_DIR):
        m = re.match(r'^havola_log_\((\d{2})-(\d{2})\.(\d{4})\)\.txt$', fname)
        if m:
            dd, mm, yyyy = m.groups()
            dates.append(f"{yyyy}-{mm}-{dd}")
    return sorted(dates, reverse=True)

def papkaoch_read_log(date_str):
    """'YYYY-MM-DD' sana uchun log fayl matnini o'qiydi (bo'lmasa — bo'sh matn).
    ESKI (butun faylni bitta matn sifatida qaytaruvchi) funksiya — hozir
    ishlatilmaydi (sahifalash uchun quyidagi papkaoch_read_log_page()ga
    o'tildi), lekin kelajakda kerak bo'lib qolishi mumkin bo'lgani uchun olib
    tashlanmadi."""
    try:
        d = datetime.strptime(date_str, '%Y-%m-%d')
    except ValueError:
        return ''
    fname = d.strftime('havola_log_(%d-%m.%Y).txt')
    fpath = os.path.join(PAPKAOCH_LOG_DIR, fname)
    if not os.path.exists(fpath):
        return ''
    try:
        with open(fpath, 'r', encoding='utf-8') as f:
            return f.read()
    except OSError:
        return ''

def papkaoch_read_log_page(date_str, page=1, page_size=200):
    """'YYYY-MM-DD' sana uchun log faylni QATOR-QATOR (sahifalab) o'qiydi.

    MUHIM: bitta kunlik log fayl 10 000-50 000 qatorgacha borishi mumkin —
    buni bitta uzun matn sifatida sahifaga chiqarish noqulay (sahifa juda
    uzun/og'ir bo'lib qoladi). Shu sabab bu yerda faylni to'liq xotiraga
    o'qib olamiz-u (bu o'zi tez, chunki matn fayl, hajmi ko'p bo'lsa ham
    bir necha MB dan oshmaydi), lekin FRONTENDGA faqat so'ralgan sahifadagi
    qatorlarni qaytaramiz — "Oxirgi faoliyat" bo'limidagi kabi ENG SO'NGGI
    yozuv birinchi (yuqorida) ko'rinadigan qilib, qatorlar teskari tartibda
    (oxiridan boshiga) sahifalanadi.

    Qaytadi: (lines: list[str], total_lines: int, total_pages: int)
    """
    try:
        d = datetime.strptime(date_str, '%Y-%m-%d')
    except ValueError:
        return [], 0, 1
    fname = d.strftime('havola_log_(%d-%m.%Y).txt')
    fpath = os.path.join(PAPKAOCH_LOG_DIR, fname)
    if not os.path.exists(fpath):
        return [], 0, 1
    try:
        with open(fpath, 'r', encoding='utf-8') as f:
            all_lines = [ln.rstrip('\n') for ln in f]
    except OSError:
        return [], 0, 1

    all_lines.reverse()  # eng so'nggi yozuv birinchi bo'lsin
    total_lines = len(all_lines)
    total_pages = max(1, math.ceil(total_lines / page_size))
    if page < 1:
        page = 1
    if page > total_pages:
        page = total_pages
    start = (page - 1) * page_size
    end = start + page_size
    return all_lines[start:end], total_lines, total_pages


_URL_RE = re.compile(r'https?://\S+', re.IGNORECASE)
_TAG_RE = re.compile(r'#([^\s#]+)', re.UNICODE)
_NUMBERED_LINE_RE = re.compile(r'^\s*(\d+)\s*[.\)]\s*(\S+)\s*$')

def parse_telegram_message(text):
    """Telegram xabar matnini tahlil qiladi.

    Kutilgan format:
        #ota_papka
        #bola_papka
        1. https://...
        2. https://...

    Qaytadi: {'tags': [...], 'links': [(raqam, url), ...]} — agar xabarda
    KAMIDA bitta teg VA kamida bitta raqamlangan havola bo'lsa. Aks holda —
    None (bu xabar botga tegishli emas, e'tiborsiz qoldiriladi — talab
    hujjatidagi "faqat # belgisi bor postlardagi URL manzillar mavjud
    xabarlarni qayta ishlaydi" qoidasiga mos)."""
    if not text:
        return None

    tags = _TAG_RE.findall(text)
    if not tags:
        return None

    links = []
    for line in text.splitlines():
        m = _NUMBERED_LINE_RE.match(line)
        if not m:
            continue
        order_num = int(m.group(1))
        candidate = m.group(2)
        url_match = _URL_RE.match(candidate)
        if url_match:
            links.append((order_num, url_match.group(0)))

    if not links:
        return None

    return {'tags': tags, 'links': links}


def match_domain(url, domains):
    """Berilgan URL'ni domenlar ro'yxati (papkaoch_domains qatorlari) bilan
    solishtiradi. Eng UZUN (eng aniq) mos kelgan base_url'ga ega domen
    qaytariladi. Hech biri mos kelmasa — None."""
    try:
        parsed = urlparse(url)
        host = (parsed.netloc or '').lower()
        host = host[4:] if host.startswith('www.') else host
    except ValueError:
        return None
    if not host:
        return None

    best = None
    best_len = -1
    for d in domains:
        try:
            dparsed = urlparse(d['base_url'])
            dhost = (dparsed.netloc or dparsed.path or '').lower().strip('/')
            dhost = dhost[4:] if dhost.startswith('www.') else dhost
        except ValueError:
            continue
        if not dhost:
            continue
        if host == dhost or host.endswith('.' + dhost):
            if len(dhost) > best_len:
                best = d
                best_len = len(dhost)
    return best


def extract_domain_label(url):
    """Bazada topilmagan domenlar uchun — URL'dan ko'rsatiladigan domen
    nomini ajratib oladi (masalan "https://namuna.uz/post/1" -> "namuna.uz")."""
    try:
        parsed = urlparse(url)
        host = (parsed.netloc or '').lower()
        host = host[4:] if host.startswith('www.') else host
        return host or url
    except ValueError:
        return url


def build_folder_name(order_num, domain_label, short_code):
    """"[raqam]. [Domen nomi] [qisqa so'z]" — masalan "1. Instagram (ins)"
    yoki bazada topilmagan domenlar uchun "2. namuna.uz (nul)"."""
    return f"{order_num}. {domain_label} ({short_code})"


def sanitize_path_component(name):
    """Papka nomidan fayl tizimida taqiqlangan belgilarni olib tashlaydi
    (Windows uchun: \\ / : * ? " < > |)."""
    return re.sub(r'[\\/:*?"<>|]', '_', name).strip()


def resolve_full_path(root_path, tags, folder_name):
    """root_path / ota_papka / [bola_papka(lar)] / folder_name — to'liq manzilni
    hosil qiladi. Teglar ichma-ich (bir nechta bo'lishi mumkin) sifatida
    ishlatiladi (talab hujjatidagi 1.2 va 2-bo'limlarga mos)."""
    parts = [sanitize_path_component(t) for t in tags] + [sanitize_path_component(folder_name)]
    return os.path.join(root_path, *parts)


def check_duplicate(conn, url, tags, order_num):
    """Bazadan (papkaoch_entries, status='created' bo'lganlar orasidan) shu
    havola/joy/raqam uchun takroriylikni tekshiradi.

    QOIDA (bitta teg-yo'l — ya'ni bitta "papka" — ICHIDA har bir tartib
    raqami FAQAT bitta havolaga tegishli bo'lishi kerak, va har bir havola
    FAQAT bitta tartib raqamiga ega bo'lishi kerak):
      1) Xuddi shu teg-yo'l (tags) ICHIDA — yoki xuddi shu TARTIB RAQAMI
         (boshqa havola bilan) bandmi, yoki xuddi shu HAVOLA (boshqa tartib
         raqami bilan) allaqachon ishlatilganmi — ikkalasi ham DARHOL rad
         etiladi (pending/HA-YO'Q so'ralmaydi, chunki bu — foydalanuvchi
         xatosi/qayta yuborishi, yangi joy tanlashga hojat yo'q).
      2) Aks holda — xuddi shu HAVOLA boshqa teg-yo'lda (boshqa papkada)
         allaqachon yaratilgan bo'lsa — bu haqiqiy "boshqa joyda takror"
         holati, shuning uchun Telegram orqali HA/YO'Q so'raladi (foydalanuvchi
         xohlasa shu yerda ham yangi papka ochishi mumkin).
      3) Hech biri topilmasa — yangi papka yaratiladi.

    Qaytadi:
      ('none', None)
      ('same_folder_conflict', row)   — 1-band (darhol rad etiladi)
      ('different_context', row)      — 2-band (HA/YO'Q so'raladi)
    """
    tags_norm = json.dumps(tags, ensure_ascii=False)

    same_folder = conn.execute(
        "SELECT * FROM papkaoch_entries WHERE tags_json=? AND status='created' "
        "AND (order_num=? OR url=?) ORDER BY id DESC",
        (tags_norm, order_num, url)
    ).fetchone()
    if same_folder:
        return ('same_folder_conflict', same_folder)

    other_folder = conn.execute(
        "SELECT * FROM papkaoch_entries WHERE url=? AND status='created' ORDER BY id DESC",
        (url,)
    ).fetchone()
    if other_folder:
        return ('different_context', other_folder)

    return ('none', None)


def create_folder_on_disk(full_path):
    """Papkani yaratadi (agar allaqachon mavjud bo'lsa — qayta yaratilmaydi,
    xato ham bermaydi). Qaytadi: (ok: bool, error_message: str|None)."""
    try:
        os.makedirs(full_path, exist_ok=True)
        return True, None
    except OSError as e:
        return False, str(e)


def process_single_link(conn, root_path, domains, chat_id, message_id, tags, order_num, url):
    """Bitta havolani to'liq qayta ishlaydi: domen aniqlash -> papka nomi ->
    takroriylikni tekshirish -> (agar kerak bo'lsa) papka yaratish.

    Qaytadi: dict —
      {'action': 'created', 'folder_name':..., 'full_path':..., ...}
      {'action': 'rejected_duplicate', ...}
      {'action': 'pending_confirm', 'token': ..., ...}   (Telegramda HA/YO'Q so'raladi)
      {'action': 'error', 'error': ...}
    """
    domain = match_domain(url, domains)
    if domain:
        domain_label = domain['name']
        short_code = domain['short_code']
    else:
        domain_label = extract_domain_label(url)
        short_code = 'nul'

    folder_name = build_folder_name(order_num, domain_label, short_code)
    full_path = resolve_full_path(root_path, tags, folder_name)

    dup_kind, dup_row = check_duplicate(conn, url, tags, order_num)
    tags_label = ' -> '.join('#' + t for t in tags)

    if dup_kind == 'same_folder_conflict':
        # Ikkala holat ham bitta teg-yo'l (papka) ICHIDA — shuning uchun
        # DARHOL rad etiladi, HA/YO'Q so'ralmaydi.
        if dup_row['order_num'] == order_num and dup_row['url'] != url:
            # Xuddi shu tartib raqami — lekin BOSHQA havola — bilan bu papka
            # ichida allaqachon papka ochilgan (masalan bir xil tartib
            # raqamli, lekin har xil post/havolalar ketma-ket tashlansa).
            note = (
                f"Bu tartib raqami (#{order_num}) bilan {tags_label} ichida allaqachon papka ochilgan. "
                f"Papka nomi: {dup_row['folder_name']}"
            )
        else:
            # Xuddi shu HAVOLA (tartib raqami bir xil yoki har xil bo'lishidan
            # qat'i nazar) shu teg-yo'l ichida allaqachon papka ochgan.
            note = f"Bu havola {tags_label} ichida allaqachon papka yaratilgan. Papka nomi: {dup_row['folder_name']}"
        return {
            'action': 'rejected_duplicate',
            'folder_name': dup_row['folder_name'],
            'full_path': dup_row['full_path'],
            'note': note,
        }

    if dup_kind == 'different_context':
        # Xuddi shu havola — lekin BOSHQA papka (teg-yo'l)da — allaqachon
        # yaratilgan. Bu yerda avtomatik rad etilmaydi — Telegram orqali
        # HA/YO'Q so'raladi (foydalanuvchi shu yerda ham yangi papka ochish-
        # ochmaslikni o'zi hal qiladi).
        old_tags = json.loads(dup_row['tags_json'])
        old_tags_label = ' -> '.join('#' + t for t in old_tags)
        note = (
            f"Bu havola boshqa papka ({old_tags_label}) ichida papka ochilgan. "
            f"Papka nomi: {dup_row['folder_name']}"
        )
        token = uuid.uuid4().hex
        now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        conn.execute(
            'INSERT INTO papkaoch_pending (token, chat_id, prompt_message_id, tags_json, order_num, url, '
            'domain_name, short_code, folder_name, full_path, created_at, resolved) '
            'VALUES (?,?,?,?,?,?,?,?,?,?,?,0)',
            (token, chat_id, None, json.dumps(tags, ensure_ascii=False), order_num, url,
             domain_label, short_code, folder_name, full_path, now)
        )
        conn.commit()
        return {
            'action': 'pending_confirm',
            'token': token,
            'note': note,
            'folder_name': folder_name,
            'full_path': full_path,
        }

    # dup_kind == 'none' -> yangi papka yaratamiz
    ok, err = create_folder_on_disk(full_path)
    now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    if ok:
        conn.execute(
            'INSERT INTO papkaoch_entries (chat_id, message_id, tags_json, order_num, url, domain_name, '
            'short_code, folder_name, full_path, status, note, created_at) '
            'VALUES (?,?,?,?,?,?,?,?,?,?,?,?)',
            (chat_id, message_id, json.dumps(tags, ensure_ascii=False), order_num, url, domain_label,
             short_code, folder_name, full_path, 'created', None, now)
        )
        conn.commit()
        return {'action': 'created', 'folder_name': folder_name, 'full_path': full_path}
    else:
        conn.execute(
            'INSERT INTO papkaoch_entries (chat_id, message_id, tags_json, order_num, url, domain_name, '
            'short_code, folder_name, full_path, status, note, created_at) '
            'VALUES (?,?,?,?,?,?,?,?,?,?,?,?)',
            (chat_id, message_id, json.dumps(tags, ensure_ascii=False), order_num, url, domain_label,
             short_code, folder_name, full_path, 'error', err, now)
        )
        conn.commit()
        return {'action': 'error', 'error': err, 'folder_name': folder_name, 'full_path': full_path}


# --- Telegram API bilan ishlash (TARMOQ TALAB QILADI — bu yerda sinalmagan,
# faqat rasmiy Telegram Bot API hujjatiga muvofiq yozilgan) ---

def tg_api_call(method, payload, timeout=15):
    """Telegram Bot API'ga so'rov yuboradi. TOKEN sozlanmagan bo'lsa — hech
    narsa qilmaydi (None qaytaradi)."""
    if not TELEGRAM_API_BASE:
        return None
    try:
        r = requests.post(f"{TELEGRAM_API_BASE}/{method}", json=payload, timeout=timeout)
        return r.json()
    except requests.RequestException as e:
        papkaoch_log([], '-', 'error', f"Telegram API xatosi ({method}): {e}")
        return None

def tg_send_message(chat_id, text, reply_markup=None):
    payload = {'chat_id': chat_id, 'text': text}
    if reply_markup:
        payload['reply_markup'] = reply_markup
    return tg_api_call('sendMessage', payload)

def tg_answer_callback(callback_query_id, text=None):
    payload = {'callback_query_id': callback_query_id}
    if text:
        payload['text'] = text
    return tg_api_call('answerCallbackQuery', payload)

def tg_edit_message_text(chat_id, message_id, text, reply_markup=None):
    payload = {'chat_id': chat_id, 'message_id': message_id, 'text': text}
    if reply_markup is not None:
        payload['reply_markup'] = reply_markup
    return tg_api_call('editMessageText', payload)

def tg_get_updates(offset=None, timeout=25):
    payload = {'timeout': timeout}
    if offset is not None:
        payload['offset'] = offset
    return tg_api_call('getUpdates', payload, timeout=timeout + 10)


def papkaoch_handle_message(update):
    """Bitta Telegram xabarini (yangi post) qayta ishlaydi: tegishli guruhdan
    kelganmi tekshiradi, xabarni tahlil qiladi, har bir havolani alohida
    process_single_link() orqali ishlaydi, lekin natijalarni GURUHLAB —
    BITTA umumiy hisobot xabari qilib guruhga qaytaradi (post ichida nechta
    link bo'lishidan qat'i nazar — har bir link uchun alohida xabar emas)."""
    msg = update.get('message') or update.get('channel_post')
    if not msg:
        return
    chat = msg.get('chat', {})
    chat_id = str(chat.get('id', ''))

    # Xavfsizlik: faqat sozlangan GURUHDAN kelgan xabarlar qayta ishlanadi
    if not TELEGRAM_GROUP_ID or chat_id != str(TELEGRAM_GROUP_ID):
        return

    text = msg.get('text') or msg.get('caption') or ''
    parsed = parse_telegram_message(text)
    if not parsed:
        return  # bu xabar botga tegishli emas (# yoki raqamlangan link yo'q)

    root_path = get_setting('papkaoch_root_path', '')
    if not root_path:
        tg_send_message(chat_id, "⚠️ Xatolik: papka manzili (root_path) hali sozlanmagan. Veb-sahifadagi \"Papkaoch\" bo'limidan sozlang.")
        return

    tags = parsed['tags']
    tags_label = ' -> '.join('#' + t for t in tags)
    message_id = msg.get('message_id')

    with get_db() as conn:
        domains = [dict(r) for r in conn.execute('SELECT * FROM papkaoch_domains').fetchall()]
        results = []
        for order_num, url in parsed['links']:
            res = process_single_link(conn, root_path, domains, chat_id, message_id, tags, order_num, url)
            results.append((order_num, url, res))

    created = [(o, u, r) for o, u, r in results if r['action'] == 'created']
    rejected = [(o, u, r) for o, u, r in results if r['action'] == 'rejected_duplicate']
    pending = [(o, u, r) for o, u, r in results if r['action'] == 'pending_confirm']
    errors = [(o, u, r) for o, u, r in results if r['action'] == 'error']

    # --- GURUHGA: bitta umumiy hisobot xabari (yaratildi/rad etildi/tasdiq kerak) ---
    lines = []
    if created:
        # Har bir yaratilgan papka nomini alohida yozish shart emas — faqat
        # umumiy son yetarli. Papka nomi faqat TAKRORIY yoki TASDIQ KERAK
        # bo'lgan hollarda (pastda) ko'rsatiladi, chunki o'sha holatlarda
        # foydalanuvchiga aynan QAYSI papka haqida gap ketayotgani muhim.
        lines.append(f"✅ \"{tags_label}\" papka ichiga {len(created)} ta yangi papka ochildi.")
    if rejected:
        if lines:
            lines.append('')
        lines.append(f"❌ {len(rejected)} ta havola takroriy bo'lgani uchun ochilmadi:")
        for o, u, r in rejected:
            lines.append(f"  • #{o}: {r['note']}")
    if pending:
        if lines:
            lines.append('')
        lines.append(f"⏳ {len(pending)} ta havola boshqa joyda allaqachon mavjud — quyidagi tugmalar orqali tasdiqlang:")
        for o, u, r in pending:
            lines.append(f"  • #{o}: {r['note']}\n     Taklif etilayotgan nom: {r['folder_name']}")

    keyboard = None
    if pending:
        keyboard = {'inline_keyboard': [
            [
                {'text': f"✅ #{o} HA", 'callback_data': f"papkaoch_yes:{r['token']}"},
                {'text': f"❌ #{o} YO'Q", 'callback_data': f"papkaoch_no:{r['token']}"},
            ]
            for o, u, r in pending
        ]}

    if lines:
        sent = tg_send_message(chat_id, '\n'.join(lines), reply_markup=keyboard)
        if pending and sent and sent.get('ok'):
            sent_message_id = sent['result']['message_id']
            with get_db() as conn:
                for o, u, r in pending:
                    conn.execute(
                        'UPDATE papkaoch_pending SET prompt_message_id=? WHERE token=?',
                        (sent_message_id, r['token'])
                    )
                conn.commit()

    # --- ADMINGA (shaxsiy chat): xatoliklar, agar bo'lsa, bitta umumiy xabar bilan ---
    if errors:
        err_lines = [f"⚠️ \"{tags_label}\" postida {len(errors)} ta xatolik yuz berdi:"]
        for o, u, r in errors:
            err_lines.append(f"  • #{o}: \"{r['folder_name']}\" — {r['error']}")
        target_chat = TELEGRAM_ADMIN_CHAT_ID or chat_id
        tg_send_message(target_chat, '\n'.join(err_lines))

    # --- LOG FAYLGA: har bir havola alohida yoziladi (bu -- audit yozuvi,
    # foydalanuvchiga ko'rinadigan Telegram xabari emas, shuning uchun bir
    # nechta qator bo'lishi "bitta hisobot" talabiga zid emas) ---
    for o, u, r in results:
        action = r['action']
        if action == 'created':
            papkaoch_log(tags, u, '+', r['folder_name'])
        elif action == 'rejected_duplicate':
            papkaoch_log(tags, u, '-', r['note'])
        elif action == 'pending_confirm':
            papkaoch_log(tags, u, '?', r['note'])
        elif action == 'error':
            papkaoch_log(tags, u, 'error', r.get('error', ''))


def papkaoch_handle_callback(update):
    """HA/YO'Q tugmasi bosilganda ishga tushadi — pending yozuvni topib,
    javobga qarab papka yaratadi yoki bekor qiladi."""
    cq = update.get('callback_query')
    if not cq:
        return
    data = cq.get('data', '')
    callback_id = cq.get('id')
    msg = cq.get('message', {})
    chat_id = str(msg.get('chat', {}).get('id', ''))

    if not data.startswith('papkaoch_yes:') and not data.startswith('papkaoch_no:'):
        return

    is_yes = data.startswith('papkaoch_yes:')
    token = data.split(':', 1)[1]

    with get_db() as conn:
        row = conn.execute('SELECT * FROM papkaoch_pending WHERE token=?', (token,)).fetchone()
        if not row or row['resolved']:
            tg_answer_callback(callback_id, "Bu so'rov allaqachon javoblangan yoki topilmadi.")
            return

        conn.execute('UPDATE papkaoch_pending SET resolved=1 WHERE token=?', (token,))

        tags = json.loads(row['tags_json'])
        if is_yes:
            ok, err = create_folder_on_disk(row['full_path'])
            now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            status = 'created' if ok else 'error'
            conn.execute(
                'INSERT INTO papkaoch_entries (chat_id, message_id, tags_json, order_num, url, domain_name, '
                'short_code, folder_name, full_path, status, note, created_at) '
                'VALUES (?,?,?,?,?,?,?,?,?,?,?,?)',
                (row['chat_id'], None, row['tags_json'], row['order_num'], row['url'], row['domain_name'],
                 row['short_code'], row['folder_name'], row['full_path'], status, err, now)
            )
            conn.commit()
            # MUHIM: asl (bir nechta tasdiq tugmasini o'z ichiga olishi mumkin bo'lgan)
            # umumiy xabarni TAHRIRLAMAYMIZ — aks holda bir tugma bosilganda boshqa
            # hali javob berilmagan tugmalar ham yo'qolib qolar edi. Shuning uchun
            # ALOHIDA, qisqa tasdiq xabari yuboramiz.
            if ok:
                tg_send_message(chat_id, f"✅ #{row['order_num']}: Yangi papka ochildi — {row['folder_name']}")
                papkaoch_log(tags, row['url'], '+', f"tasdiqlandi -> {row['folder_name']}")
            else:
                tg_send_message(chat_id, f"⚠️ #{row['order_num']}: \"{row['folder_name']}\" papkasini yaratib bo'lmadi — {err}")
                papkaoch_log(tags, row['url'], 'error', err or '')
            tg_answer_callback(callback_id, "Papka ochildi." if ok else "Xatolik yuz berdi.")
        else:
            conn.commit()
            tg_send_message(chat_id, f"❌ #{row['order_num']}: Bekor qilindi — yangi papka ochilmadi ({row['folder_name']}).")
            papkaoch_log(tags, row['url'], '-', "foydalanuvchi rad etdi")
            tg_answer_callback(callback_id, "Bekor qilindi.")


_papkaoch_bot_thread = None
_papkaoch_bot_stop = threading.Event()
_papkaoch_bot_last_error = None   # oxirgi getUpdates/loop xatosi (bo'lmasa None) -- "Bot holati"da ko'rsatiladi
_papkaoch_bot_last_poll_ok = None  # oxirgi so'rov muvaffaqiyatli o'tdimi

def papkaoch_bot_loop():
    """Fon jarayoni (background thread): Telegram getUpdates orqali doimiy
    ravishda yangi xabarlarni so'rab turadi (long polling). Faqat TOKEN va
    GURUH ID .env faylida sozlangan bo'lsa ishga tushadi."""
    global _papkaoch_bot_last_error, _papkaoch_bot_last_poll_ok
    if not papkaoch_bot_configured():
        return

    # MUHIM: agar shu bot TOKENi uchun avval (masalan boshqa dastur yoki
    # sinov paytida) WEBHOOK sozlangan bo'lsa, Telegram getUpdates orqali
    # HECH QACHON yangi xabar qaytarmaydi — va bu xato hech qanday belgisiz,
    # "jim" sodir bo'ladi (aynan shu holat "bot ishlab turibdi" ko'rinib,
    # lekin hech narsa qilmasligi mumkin edi). Shu sabab bot ishga tushganda
    # avval eski webhookni tozalab tashlaymiz.
    wh = tg_api_call('deleteWebhook', {'drop_pending_updates': False})
    if not (wh and wh.get('ok')):
        papkaoch_log([], '-', 'error', f"deleteWebhook chaqiruvi muvaffaqiyatsiz bo'ldi: {wh}")

    offset = None
    last_logged_error = None
    while not _papkaoch_bot_stop.is_set():
        try:
            data = tg_get_updates(offset=offset, timeout=25)

            if not data:
                # tg_get_updates() allaqachon tarmoq xatosini o'zi log qiladi
                # (tg_api_call ichida) — bu yerda faqat holatni belgilaymiz
                _papkaoch_bot_last_poll_ok = False
                _papkaoch_bot_last_error = "Telegram serveridan javob kelmadi (tarmoq xatosi)"
                time_module.sleep(3)
                continue

            if not data.get('ok'):
                # MUHIM: bu — avval butunlay jim qoldirilgan holat edi. Endi
                # sababini (Telegramning o'z xato matni bilan) log faylga
                # yozamiz — lekin har 3 soniyada QAYTA-QAYTA yozib log faylni
                # to'ldirib yubormaslik uchun, xato o'zgarmagunicha faqat
                # BIR MARTA yoziladi.
                err_desc = f"[{data.get('error_code')}] {data.get('description')}"
                _papkaoch_bot_last_poll_ok = False
                _papkaoch_bot_last_error = err_desc
                if last_logged_error != err_desc:
                    papkaoch_log([], '-', 'error', f"getUpdates xatosi: {err_desc}")
                    last_logged_error = err_desc
                time_module.sleep(3)
                continue

            _papkaoch_bot_last_poll_ok = True
            _papkaoch_bot_last_error = None
            last_logged_error = None

            for update in data.get('result', []):
                offset = update['update_id'] + 1
                try:
                    if 'callback_query' in update:
                        papkaoch_handle_callback(update)
                    else:
                        papkaoch_handle_message(update)
                except Exception as e:
                    papkaoch_log([], '-', 'error', f"Update qayta ishlashda xato: {e}")
        except Exception as e:
            _papkaoch_bot_last_poll_ok = False
            _papkaoch_bot_last_error = str(e)
            papkaoch_log([], '-', 'error', f"Bot loop xatosi: {e}")
            time_module.sleep(5)

def start_papkaoch_bot():
    """Fon oqimini (agar hali ishga tushmagan bo'lsa) ishga tushiradi."""
    global _papkaoch_bot_thread
    if not papkaoch_bot_configured():
        return
    if _papkaoch_bot_thread and _papkaoch_bot_thread.is_alive():
        return
    _papkaoch_bot_thread = threading.Thread(target=papkaoch_bot_loop, daemon=True)
    _papkaoch_bot_thread.start()


def papkaoch_run_validation():
    """"Tekshiruv" rejimi: bazadagi ('created' statusli) har bir yozuv uchun
    tegishli papka hali ham diskda mavjudligini tekshiradi. Natijalar ro'yxati
    (muammoli qatorlar) qaytariladi — web sahifada ko'rsatiladi."""
    issues = []
    with get_db() as conn:
        rows = conn.execute("SELECT * FROM papkaoch_entries WHERE status='created' ORDER BY id DESC").fetchall()
    for row in rows:
        if not row['full_path']:
            continue
        if not os.path.isdir(row['full_path']):
            issues.append({
                'id': row['id'],
                'url': row['url'],
                'folder_name': row['folder_name'],
                'full_path': row['full_path'],
                'problem': "Papka diskda topilmadi (o'chirilgan yoki ko'chirilgan bo'lishi mumkin)",
            })
    return issues


# "sana"/"edited_at" ustunlaridagi qiymatlar har xil uzunlikda ("YYYY-MM-DD" — 10,
# "YYYY-MM-DD HH:MM" — 16, yoki "YYYY-MM-DD HH:MM:SS" — 19 belgi) saqlanishi mumkin.
# Sana filtrida (build_filters) to'g'ri solishtirish uchun ularni SQL darajasida
# bir xil uzunlikka ("YYYY-MM-DD HH:MM:SS") keltiramiz — soat berilmagan qatorlarga
# kun boshi (00:00:00) sifatida qaraymiz.
_DT_NORMALIZE_SQL = """
    CASE length(COALESCE(edited_at, sana))
        WHEN 10 THEN COALESCE(edited_at, sana) || ' 00:00:00'
        WHEN 16 THEN COALESCE(edited_at, sana) || ':00'
        ELSE COALESCE(edited_at, sana)
    END
"""

def _normalize_dt(val, end_of_day=False):
    """Filtr sifatida kiritilgan sana/sana+soat qiymatini "YYYY-MM-DD HH:MM:SS"
    formatiga keltiradi. Faqat sana berilgan bo'lsa (10 belgi), end_of_day=True
    bo'lsa kun oxiri (23:59:59), aks holda kun boshi (00:00:00) qo'shiladi."""
    if not val:
        return val
    if len(val) == 10:
        return val + (' 23:59:59' if end_of_day else ' 00:00:00')
    if len(val) == 16:
        return val + ':00'
    return val

def build_filters(args):
    """So'rov argumentlaridan umumiy WHERE shartini va parametrlarni yasaydi."""
    board = get_board(args)
    clause = ' AND board = ?'
    params = [board]

    status = args.get('status', '')
    date_from = args.get('date_from', '')
    date_to = args.get('date_to', '')
    likes_min = args.get('likes_min', '')
    likes_max = args.get('likes_max', '')
    search = args.get('search', '')

    if status:
        clause += ' AND status = ?'
        params.append(status)
    if date_from:
        # MUHIM: bazadagi "sana"/"edited_at" ustunlari har xil aniqlikda saqlanishi
        # mumkin — ko'p qatorlarda (ayniqsa hali tahrirlanmagan, "Papkadan import"
        # orqali kiritilgan qatorlarda) faqat sana bor, soat umuman saqlanmaydi.
        #
        # Agar bunday soatsiz qatorni har doim "kun boshi — 00:00:00" deb hisoblab,
        # to'g'ridan-to'g'ri filtr soati bilan solishtirsak — filtrda tasodifan
        # (masalan brauzer joriy vaqtni avtomatik to'ldirgani sabab) "00:00" dan
        # KATTA biror soat (masalan "08:26") turgan bo'lsa, o'sha kunning o'ziga
        # tegishli, lekin soati noma'lum qatorlar NOTO'G'RI ravishda filtrdan
        # chiqib ketadi — garchi ular aynan izlangan kunga tegishli bo'lsa ham.
        #
        # Shu sabab qoida: agar qatorning SOATI umuman saqlanmagan bo'lsa va
        # uning KUNI filtr chegarasidagi kun bilan mos kelsa — soatidan qat'i
        # nazar HAR DOIM ko'rsatiladi ("shubhadan foyda" tamoyili). Aniq soat
        # bo'yicha solishtirish esa faqat HAQIQIY vaqti saqlangan qatorlarga
        # (masalan tahrirlangan yozuvlarga) qo'llaniladi — bu aynan smena
        # oralig'ida filtrlashni ("2026-08-03 08:00" dan "2026-08-04 07:30"
        # gacha) to'g'ri ishlatish uchun kerak.
        df_full = _normalize_dt(date_from, end_of_day=False)
        df_date = df_full[:10]
        clause += f'''
            AND (
                substr(COALESCE(edited_at, sana), 1, 10) > ?
                OR (
                    substr(COALESCE(edited_at, sana), 1, 10) = ?
                    AND (
                        length(COALESCE(edited_at, sana)) = 10
                        OR ({_DT_NORMALIZE_SQL}) >= ?
                    )
                )
            )'''
        params.extend([df_date, df_date, df_full])
    if date_to:
        # Xuddi shu tamoyil — oxirgi chegara kuniga tegishli, soati noma'lum
        # qatorlar ham har doim ko'rsatiladi; aniq soat faqat haqiqiy vaqti
        # saqlangan qatorlarga qo'llaniladi.
        dt_full = _normalize_dt(date_to, end_of_day=True)
        dt_date = dt_full[:10]
        clause += f'''
            AND (
                substr(COALESCE(edited_at, sana), 1, 10) < ?
                OR (
                    substr(COALESCE(edited_at, sana), 1, 10) = ?
                    AND (
                        length(COALESCE(edited_at, sana)) = 10
                        OR ({_DT_NORMALIZE_SQL}) <= ?
                    )
                )
            )'''
        params.extend([dt_date, dt_date, dt_full])
    if likes_min:
        count_col = COUNT_COLUMN.get(board, 'likes')
        clause += f' AND {count_col} >= ?'
        params.append(int(likes_min))
    if likes_max:
        count_col = COUNT_COLUMN.get(board, 'likes')
        clause += f' AND {count_col} <= ?'
        params.append(int(likes_max))
    if search:
        clause += ' AND url LIKE ?'
        params.append(f'%{search}%')

    return clause, params

SORT_COLUMNS = {
    'id': 'id',
    'sana': 'COALESCE(edited_at, sana)',
    'likes': 'likes',
    'status': 'status',
    'umumiy_izohlar': 'umumiy_izohlar',
    'ajratilgan_izohlar': 'ajratilgan_izohlar',
    # Yangi kiritilgan YOKI oxirgi marta tahrirlangan vaqt — qaysi biri kechroq
    # bo'lsa o'sha olinadi. Veb sahifada "eng so'nggi amal tepada" tartibini
    # ta'minlash uchun standart saralash mezoni sifatida ishlatiladi.
    'last_activity': 'COALESCE(edited_at, created_at)',
}

def format_sana(row):
    """'Sana' ustuni uchun matn: agar qator tahrirlangan bo'lsa 'dastlabki / oxirgi tahrir sanasi',
    aks holda faqat dastlabki sana. Ko'rinishda faqat sana (soatsiz) chiqariladi."""
    sana = row['sana'] or ''
    edited = row['edited_at'] if 'edited_at' in row.keys() else None
    if edited:
        edited_date = edited[:10]  # faqat "YYYY-MM-DD" qismi
        return f"{sana} / {edited_date}"
    return sana

@app.route('/')
def index():
    return send_from_directory('static', 'index.html')

# --- Admin login/logout/holat ---
@app.route('/api/admin/login', methods=['POST'])
def admin_login():
    data = request.json or {}
    password = data.get('password', '')
    if password and password == ADMIN_PASSWORD:
        session.permanent = True
        session['is_admin'] = True
        return jsonify({'ok': True})
    return jsonify({'ok': False, 'error': "Parol noto'g'ri"}), 401

@app.route('/api/admin/logout', methods=['POST'])
def admin_logout():
    session.pop('is_admin', None)
    return jsonify({'ok': True})

@app.route('/api/admin/status', methods=['GET'])
def admin_status():
    return jsonify({'is_admin': bool(session.get('is_admin'))})

# --- Havolaning kunlik tartib raqami uchun taklif ---
@app.route('/api/statuses', methods=['GET'])
def get_statuses():
    """Shu board uchun status ro'yxatini qaytaradi. T.hisobot uchun — bu
    ro'yxat OCHIQ: oldindan belgilangan 6 ta ("Umumiy", "Tashrif", ...) ga
    qo'shimcha, bazada haqiqatan mavjud bo'lgan boshqa (papka tuzilishidan
    kelib chiqqan, masalan "Таълим") statuslar ham qo'shiladi — shunda
    "+ Yangi havola" va filtr ro'yxatlarida ular ham ko'rinadi."""
    board = get_board(request.args)
    known = list(STATUS_OPTIONS.get(board, STATUS_OPTIONS[DEFAULT_BOARD]))
    if board != 't_hisobot':
        return jsonify({'statuses': known})
    with get_db() as conn:
        rows = conn.execute(
            'SELECT DISTINCT status FROM links WHERE board=? AND status IS NOT NULL AND status != ""',
            (board,)
        ).fetchall()
    discovered = sorted({r['status'] for r in rows} - set(known))
    return jsonify({'statuses': known + discovered})

@app.route('/api/next-order-num', methods=['GET'])
def next_order_num():
    board = get_board(request.args)
    sana = request.args.get('sana', '')
    status = request.args.get('status', '')
    if not sana:
        return jsonify({'next': 1})
    with get_db() as conn:
        if board == 't_hisobot':
            # T.hisobot'da tartib raqami har bir STATUS (Umumiy/Tashrif/...) ichida
            # ALOHIDA-ALOHIDA hisoblanadi — chunki tarmoq papkasida ham har bir rukun
            # papkasi o'zining "1, 2, 3..." raqamlashuvidan boshlanadi. Shu sabab
            # bu yerda board+sana bilan birga status ham hisobga olinishi shart —
            # aks holda taklif qilingan raqam noto'g'ri (boshqa statusning
            # raqamlari bilan qo'shib hisoblangan) bo'lib chiqadi.
            row = conn.execute(
                'SELECT COALESCE(MAX(order_num), 0) as m FROM links WHERE board=? AND sana=? AND status=?',
                (board, sana, status)
            ).fetchone()
        else:
            row = conn.execute(
                'SELECT COALESCE(MAX(order_num), 0) as m FROM links WHERE board=? AND sana=?',
                (board, sana)
            ).fetchone()
    return jsonify({'next': (row['m'] or 0) + 1})

@app.route('/api/last-folder-path', methods=['GET'])
def last_folder_path():
    # MUHIM: "Papkadan import qilish" oynasi ochilganda "Papka manzili" maydoni
    # DEFAULT holatda har doim "Papkaoch" bo'limida sozlangan UMUMIY manzilni
    # ko'rsatishi kerak (jaloba/fishing/t_hisobot — barchasi uchun bir xil) —
    # avval qo'lda kiritib qo'yilgan (board'ga xos) manzil emas. Agar
    # foydalanuvchi shu oynada boshqa manzilni qo'lda kiritsa — bu FAQAT o'sha
    # safargi ko'rish uchun ishlatiladi (preview/apply so'rovlarida yuboriladi),
    # lekin keyingi safar oyna qayta ochilganda yana "Papkaoch"dagi umumiy
    # manzil ko'rsatiladi (board bo'yicha alohida "eslab qolish" endi yo'q).
    return jsonify({'folder_path': get_setting('papkaoch_root_path', '') or ''})

# ============================================================================
# "PAPKAOCH" BO'LIMI — API endpointlari
# ============================================================================

@app.route('/api/papkaoch/status', methods=['GET'])
def papkaoch_status():
    """Bot sozlanganmi, ishga tushirilganmi va qaysi ma'lumotlar .env orqali
    kelayotganini (qiymatlarini OCHMASDAN) ko'rsatadi."""
    return jsonify({
        'token_configured': bool(TELEGRAM_BOT_TOKEN),
        'group_id_configured': bool(TELEGRAM_GROUP_ID),
        'admin_chat_configured': bool(TELEGRAM_ADMIN_CHAT_ID),
        'bot_running': bool(_papkaoch_bot_thread and _papkaoch_bot_thread.is_alive()),
        'last_poll_ok': _papkaoch_bot_last_poll_ok,
        'last_error': _papkaoch_bot_last_error,
        'root_path': get_setting('papkaoch_root_path', '') or '',
    })

@app.route('/api/papkaoch/settings', methods=['POST'])
def papkaoch_update_settings():
    """Faqat "root_path" (papka manzili) shu yerdan sozlanadi — Bot TOKEN va
    Guruh ID xavfsizlik uchun FAQAT serverdagi .env faylida saqlanadi va
    veb-interfeys orqali o'zgartirilmaydi. Bu sozlama — jaloba/fishing/t_hisobot
    bo'limlaridagi "papka manzili" kabi — istalgan foydalanuvchi kirita oladi,
    faqat admin emas."""
    data = request.json or {}
    root_path = (data.get('root_path') or '').strip()
    if not root_path:
        return jsonify({'error': "Papka manzili bo'sh bo'lishi mumkin emas"}), 400
    set_setting('papkaoch_root_path', root_path)
    return jsonify({'ok': True, 'root_path': root_path})

@app.route('/api/papkaoch/domains', methods=['GET'])
def papkaoch_list_domains():
    with get_db() as conn:
        rows = conn.execute('SELECT * FROM papkaoch_domains ORDER BY name').fetchall()
    return jsonify({'items': [dict(r) for r in rows]})

@app.route('/api/papkaoch/domains', methods=['POST'])
@admin_required
def papkaoch_add_domain():
    data = request.json or {}
    name = (data.get('name') or '').strip()
    base_url = (data.get('base_url') or '').strip()
    short_code = (data.get('short_code') or '').strip()
    if not name or not base_url or not short_code:
        return jsonify({'error': "Nomi, asosiy havola va qisqa so'z to'ldirilishi shart"}), 400
    now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    with get_db() as conn:
        cur = conn.execute(
            'INSERT INTO papkaoch_domains (name, base_url, short_code, created_at) VALUES (?,?,?,?)',
            (name, base_url, short_code, now)
        )
        conn.commit()
        new_id = cur.lastrowid
        row = conn.execute('SELECT * FROM papkaoch_domains WHERE id=?', (new_id,)).fetchone()
    return jsonify(dict(row))

@app.route('/api/papkaoch/domains/<int:did>', methods=['PUT'])
@admin_required
def papkaoch_update_domain(did):
    data = request.json or {}
    name = (data.get('name') or '').strip()
    base_url = (data.get('base_url') or '').strip()
    short_code = (data.get('short_code') or '').strip()
    if not name or not base_url or not short_code:
        return jsonify({'error': "Nomi, asosiy havola va qisqa so'z to'ldirilishi shart"}), 400
    with get_db() as conn:
        existing = conn.execute('SELECT * FROM papkaoch_domains WHERE id=?', (did,)).fetchone()
        if not existing:
            return jsonify({'error': 'Domen topilmadi'}), 404
        conn.execute(
            'UPDATE papkaoch_domains SET name=?, base_url=?, short_code=? WHERE id=?',
            (name, base_url, short_code, did)
        )
        conn.commit()
        row = conn.execute('SELECT * FROM papkaoch_domains WHERE id=?', (did,)).fetchone()
    return jsonify(dict(row))

@app.route('/api/papkaoch/domains/<int:did>', methods=['DELETE'])
@admin_required
def papkaoch_delete_domain(did):
    with get_db() as conn:
        existing = conn.execute('SELECT * FROM papkaoch_domains WHERE id=?', (did,)).fetchone()
        if not existing:
            return jsonify({'error': 'Domen topilmadi'}), 404
        conn.execute('DELETE FROM papkaoch_domains WHERE id=?', (did,))
        conn.commit()
    return jsonify({'ok': True})

@app.route('/api/papkaoch/entries', methods=['GET'])
def papkaoch_list_entries():
    """Oxirgi faoliyat — bot tomonidan qayta ishlangan havolalar ro'yxati
    (eng so'nggisi birinchi), SAHIFALANGAN holda.

    MUHIM: bu jadval vaqt o'tishi bilan 10 000-50 000 qatorgacha o'sib
    ketishi mumkin — shu sabab endi bitta katta ro'yxat o'rniga
    "page"/"per_page" bo'yicha sahifalab qaytariladi (frontendda "Oldingi/
    Keyingi" tugmalari bilan ko'rsatiladi), sahifa juda uzun bo'lib
    ketmasligi uchun."""
    try:
        page = int(request.args.get('page', 1))
    except ValueError:
        page = 1
    try:
        per_page = int(request.args.get('per_page', 50))
    except ValueError:
        per_page = 50
    per_page = max(10, min(per_page, 200))
    if page < 1:
        page = 1

    with get_db() as conn:
        total = conn.execute('SELECT COUNT(*) FROM papkaoch_entries').fetchone()[0]
        total_pages = max(1, math.ceil(total / per_page))
        if page > total_pages:
            page = total_pages
        offset = (page - 1) * per_page
        rows = conn.execute(
            'SELECT * FROM papkaoch_entries ORDER BY id DESC LIMIT ? OFFSET ?',
            (per_page, offset)
        ).fetchall()
    items = []
    for r in rows:
        d = dict(r)
        try:
            d['tags'] = json.loads(d.get('tags_json') or '[]')
        except (ValueError, TypeError):
            d['tags'] = []
        items.append(d)
    return jsonify({
        'items': items,
        'total': total,
        'page': page,
        'per_page': per_page,
        'total_pages': total_pages,
    })

@app.route('/api/papkaoch/logs/dates', methods=['GET'])
def papkaoch_log_dates():
    return jsonify({'dates': papkaoch_list_log_dates()})

@app.route('/api/papkaoch/logs', methods=['GET'])
def papkaoch_logs():
    """Bitta kunlik log faylni SAHIFALAB qaytaradi (qarang: papkaoch_read_log_page).
    MUHIM: kunlik log fayl 10 000-50 000 qatorgacha borishi mumkin — shu
    sabab endi butun matn emas, faqat so'ralgan sahifadagi qatorlar
    qaytariladi (eng so'nggi yozuv birinchi sahifada, yuqorida)."""
    date_str = request.args.get('date', '')
    try:
        page = int(request.args.get('page', 1))
    except ValueError:
        page = 1
    try:
        page_size = int(request.args.get('page_size', 200))
    except ValueError:
        page_size = 200
    page_size = max(20, min(page_size, 1000))

    if not date_str:
        return jsonify({'date': date_str, 'lines': [], 'total_lines': 0, 'page': 1, 'total_pages': 1, 'page_size': page_size})

    lines, total_lines, total_pages = papkaoch_read_log_page(date_str, page, page_size)
    return jsonify({
        'date': date_str,
        'lines': lines,
        'total_lines': total_lines,
        'page': min(max(page, 1), total_pages),
        'total_pages': total_pages,
        'page_size': page_size,
    })

@app.route('/api/papkaoch/validate', methods=['POST'])
@admin_required
def papkaoch_validate():
    issues = papkaoch_run_validation()
    return jsonify({'ok': True, 'issue_count': len(issues), 'issues': issues})

def _scan_numbered_folders(container_path):
    """Berilgan papka ichidagi raqamli pastki papkalarni (masalan '1. t.me (tm)')
    o'qib, {order_num: {'folder_name', 'main', 'comment'}} lug'atini qaytaradi."""
    folder_counts = {}
    try:
        entries = sorted(os.listdir(container_path))
    except OSError:
        return folder_counts
    for name in entries:
        full = os.path.join(container_path, name)
        if not os.path.isdir(full):
            continue
        order_num = _extract_link_folder_order(name)
        if order_num is None:
            continue
        main_count, comment_count = _count_link_folder_screenshots(full)
        folder_counts[order_num] = {'folder_name': name, 'main': main_count, 'comment': comment_count}
    return folder_counts

def _match_folder_links(folder_counts, db_links, use_category):
    """"+ Yangi havola" orqali qo'lda kiritilgan linklar ko'pincha haqiqiy papka
    tartib raqamiga emas, shunchaki "kunning navbatdagi raqami"ga ega bo'ladi.
    Shu sabab keyinroq shu kunlik papka skanerlanganda, ular yoki umuman hech
    qanday papkaga (order_num) mos kelmay "Yangi" bo'lib qolishi, YOKI battar —
    tasodifan BOSHQA (noto'g'ri) papka bilan bir xil order_num'ga tushib
    "Nomuvofiqlik" bo'lib chiqishi mumkin (masalan Facebook linki tasodifan
    "1. t.me" papkasi bilan bir xil raqamda bo'lib qolsa).

    Ikki bosqichda ishlaydi:
      1) Aniq mos kelishlarni (bir xil order_num VA — jaloba/fishing uchun —
         bir xil ijtimoiy tarmoq turi) "qulflaydi".
      2) Qolgan bo'sh papka pozitsiyalari (gap) va "erkin qolgan" (qulflanmagan)
         bazadagi yozuvlar orasidan ISHONCHLI (bir ma'noli) juftlikni qidiradi:
         - jaloba/fishing (use_category=True): papka va havola bir xil ijtimoiy
           tarmoq turiga tegishli bo'lsa VA shu turdagi yagona nomzod bo'lsa.
         - t_hisobot (use_category=False): kategoriya tushunchasi bu yerda
           ma'noga ega emas (papka — izoh QAYERDA qoldirilgani, link — MAQOLA
           manzili), shuning uchun faqat "aniq bitta bo'sh joy — aniq bitta
           erkin yozuv" holatida moslashtiradi — noaniq holatlarda hech narsa
           taxmin qilinmaydi.

    Qaytadi: (locked, realigned, used_ids)
      locked:    {order_num: link_row}   — to'g'ridan-to'g'ri (1-bosqich) mos kelganlar
      realigned: {order_num: link_row}   — moslashtirilgan (2-bosqich) juftliklar
      used_ids:  {link_id, ...}          — yuqoridagi ikkalasida ishlatilgan barcha ID'lar
    """
    locked = {}
    used_ids = set()
    for order_num, fc in folder_counts.items():
        link = db_links.get(order_num)
        if link is None:
            continue
        if use_category:
            folder_cat = guess_folder_category(fc['folder_name'])
            link_cat = guess_url_category(link['url'])
            if folder_cat and link_cat and folder_cat != link_cat:
                continue  # mos kelmaydi — bu yerda "qulflamaymiz", ehtimol boshqa joyga tegishli
        locked[order_num] = link
        used_ids.add(link['id'])

    gaps = [order_num for order_num in folder_counts if order_num not in locked]
    free_links = [r for r in db_links.values() if r['id'] not in used_ids]

    realigned = {}
    if use_category:
        newly_used = set()
        for order_num in gaps:
            folder_cat = guess_folder_category(folder_counts[order_num]['folder_name'])
            if not folder_cat:
                continue
            candidates = [r for r in free_links if r['id'] not in newly_used and guess_url_category(r['url']) == folder_cat]
            if len(candidates) == 1:
                realigned[order_num] = candidates[0]
                newly_used.add(candidates[0]['id'])
        used_ids |= newly_used
    else:
        if len(gaps) == 1 and len(free_links) == 1:
            realigned[gaps[0]] = free_links[0]
            used_ids.add(free_links[0]['id'])

    return locked, realigned, used_ids

def _build_jaloba_fishing_section(board, folder_path, sana):
    """Jaloba/Fishing uchun bitta (statussiz) bo'lim: papkadagi raqamli
    pastki papkalarni bazadagi mos linklar bilan solishtirib qatorlar hosil qiladi."""
    folder_counts = _scan_numbered_folders(folder_path)

    db_links = {}
    if sana:
        with get_db() as conn:
            rows = conn.execute(
                'SELECT * FROM links WHERE board=? AND sana=? AND order_num IS NOT NULL',
                (board, sana)
            ).fetchall()
        for r in rows:
            db_links[r['order_num']] = dict(r)

    all_orders = sorted(set(folder_counts.keys()) | set(db_links.keys()))
    locked, realigned, used_ids = _match_folder_links(folder_counts, db_links, use_category=True)

    rows_out = []
    for order_num in all_orders:
        fc = folder_counts.get(order_num)
        status_hint = None
        if order_num in locked:
            link = locked[order_num]
            status_hint = 'direct'
        elif order_num in realigned:
            link = realigned[order_num]
            status_hint = 'realigned'
        else:
            raw_link = db_links.get(order_num)
            if raw_link and raw_link['id'] in used_ids:
                # Bu yozuv boshqa (to'g'ri) papka pozitsiyasiga moslashtirildi —
                # o'zining eski (noto'g'ri) tartib raqami ostida qayta ko'rsatilmaydi.
                continue
            link = raw_link

        folder_category = guess_folder_category(fc['folder_name']) if fc else None
        link_category = guess_url_category(link['url']) if link else None
        category_mismatch = bool(
            folder_category and link_category and folder_category != link_category
            and status_hint != 'realigned'
        )

        if status_hint == 'realigned':
            match_status = 'realigned'
        elif fc and link and not category_mismatch:
            match_status = 'ok'
        elif category_mismatch:
            match_status = 'mismatch'
        elif fc and not link:
            match_status = 'no_link'
        else:
            match_status = 'no_folder'

        rows_out.append({
            'order_num': order_num,
            'folder_name': fc['folder_name'] if fc else None,
            'computed_count': fc['main'] if fc else None,
            'comment_count': fc['comment'] if fc else None,
            'link_id': link['id'] if link else None,
            'url': link['url'] if link else None,
            'current_likes': link['likes'] if link else None,
            'category_mismatch': category_mismatch,
            'status': match_status,
        })

    return [{'section_status': None, 'section_label': None, 'rows': rows_out}]

def _build_t_hisobot_sections(folder_path, sana):
    """T.hisobot uchun 'Коммент' papkasini o'qiydi. Ikki variant bo'lishi mumkin:
    1) 'Коммент' ichida to'g'ridan-to'g'ri raqamli papkalar -> bitta 'Umumiy' bo'lim.
    2) 'Коммент' ichida (istalgan chuqurlikda, istalgan nomda) ichma-ich rukun
       papkalari -> har biri uchun alohida bo'lim, "barg" (eng ichkarida,
       raqamli papkalarni o'z ichiga olgan) papkalar o'qiladi."""
    containers = _discover_t_hisobot_category_dirs(folder_path)
    if not containers:
        # Hech qanday rukun papkasi (va hatto raqamli papka ham) topilmadi —
        # "Коммент" ichidagi raqamli papkalarni to'g'ridan-to'g'ri standart
        # "Umumiy" rukuniga joylashtiramiz (masalan bo'sh papka holati)
        containers = [('Umumiy', folder_path)]

    sections = []
    for status, container_path in containers:
        folder_counts = _scan_numbered_folders(container_path)

        db_links = {}
        if sana:
            with get_db() as conn:
                rows = conn.execute(
                    'SELECT * FROM links WHERE board=? AND sana=? AND status=? AND order_num IS NOT NULL',
                    ('t_hisobot', sana, status)
                ).fetchall()
            for r in rows:
                db_links[r['order_num']] = dict(r)

        all_orders = sorted(set(folder_counts.keys()) | set(db_links.keys()))
        locked, realigned, used_ids = _match_folder_links(folder_counts, db_links, use_category=False)

        rows_out = []
        for order_num in all_orders:
            fc = folder_counts.get(order_num)
            status_hint = None
            if order_num in locked:
                link = locked[order_num]
                status_hint = 'direct'
            elif order_num in realigned:
                link = realigned[order_num]
                status_hint = 'realigned'
            else:
                raw_link = db_links.get(order_num)
                if raw_link and raw_link['id'] in used_ids:
                    continue
                link = raw_link

            # DIQQAT: T.hisobot'da papka belgisi ("(fb)") izoh QAYERDA qoldirilganini
            # bildiradi, link esa izoh qoldirilgan MAQOLA/POST manzili — bular boshqa-
            # boshqa narsa (Jaloba/Fishing'dagidek "bir xil platforma" talab qilinmaydi).
            # Shu sabab bu yerda kategoriya-mos-kelmaslik tekshiruvi qo'llanilmaydi.
            category_mismatch = False

            if status_hint == 'realigned':
                match_status = 'realigned'
            elif fc and link:
                match_status = 'ok'
            elif fc and not link:
                match_status = 'no_link'
            else:
                match_status = 'no_folder'

            rows_out.append({
                'order_num': order_num,
                'folder_name': fc['folder_name'] if fc else None,
                'computed_count': fc['main'] if fc else None,        # -> Umumiy izohlar soni taklifi
                'comment_count': fc['comment'] if fc else None,      # -> Yozilgan izohlar soni taklifi
                'link_id': link['id'] if link else None,
                'url': link['url'] if link else None,
                'current_umumiy': link['umumiy_izohlar'] if link else None,
                'current_ajratilgan': link['ajratilgan_izohlar'] if link else None,
                'category_mismatch': category_mismatch,
                'status': match_status,
            })

        sections.append({'section_status': status, 'section_label': status, 'rows': rows_out})

    return sections

def _find_board_daily_subfolder(root_path, board):
    """"Papkaoch" bo'limida sozlangan umumiy (kun.oy.yil) manzili ichidan shu
    board'ga tegishli maxsus pastki papkani qidiradi — nomi lotin/kirill,
    katta/kichik harflar bilan yozilgan bo'lishidan qat'i nazar (masalan
    "Жалоба", "жалоба", "ЖАЛОБА", "Jaloba" — barchasi bir xil deb tan olinadi).
    Topilsa to'liq yo'lini, topilmasa None qaytaradi."""
    if board == 't_hisobot':
        expected = COMMENT_FOLDER_NAMES
    else:
        expected = BOARD_SPECIAL_FOLDER_NAMES.get(board, ())
    if not expected:
        return None
    expected_lower = {e.lower() for e in expected}
    try:
        for entry in os.listdir(root_path):
            full = os.path.join(root_path, entry)
            if os.path.isdir(full) and entry.lower() in expected_lower:
                return full
    except OSError:
        return None
    return None

# --- Papkadan hisoblash: 1) KO'RIB CHIQISH (hech narsa saqlanmaydi) ---
@app.route('/api/screenshot-folder/preview', methods=['POST'])
def preview_screenshot_folder():
    data = request.json or {}
    board = data.get('board', '')
    folder_path = (data.get('folder_path') or '').strip()
    sana_override = (data.get('sana_override') or '').strip()

    if board not in ('jaloba', 'fishing', 't_hisobot'):
        return jsonify({'error': "Noto'g'ri loyiha"}), 400

    # "Papka manzili" endi IXTIYORIY: bo'sh qoldirilsa, "Papkaoch" bo'limida
    # sozlangan umumiy (kun.oy.yil) manzil ICHIDAN shu board'ga tegishli
    # maxsus pastki papka ("Жалоба"/"Фишинг"/"Коммент", lotin/kirill,
    # katta/kichik harf farqisiz) avtomatik qidirib topiladi va o'sha
    # ishlatiladi — umumiy manzilning o'zi emas (chunki u faqat kun papkasi,
    # ichida uchala bo'lim uchun alohida pastki papkalar bo'ladi).
    used_default_root = False
    if not folder_path:
        base_root = get_setting('papkaoch_root_path', '')
        if not base_root:
            return jsonify({'error': (
                "Papka manzilini kiriting, yoki avval \"Papkaoch\" bo'limida "
                "umumiy papka manzilini sozlang — shunda bu yerda bo'sh qoldirish mumkin bo'ladi."
            )}), 400
        subfolder = _find_board_daily_subfolder(base_root, board)
        if not subfolder:
            expected = COMMENT_FOLDER_NAMES if board == 't_hisobot' else BOARD_SPECIAL_FOLDER_NAMES.get(board, ())
            expected_display = '/'.join(dict.fromkeys(expected))
            return jsonify({'error': (
                f"\"Papkaoch\" bo'limida sozlangan manzil (\"{base_root}\") ichida "
                f"\"{expected_display}\" nomli pastki papka topilmadi. Shu bo'lim uchun "
                f"to'liq manzilni alohida kiriting, yoki tegishli papkani yarating."
            )}), 400
        folder_path = subfolder
        used_default_root = True

    if not os.path.isdir(folder_path):
        if used_default_root:
            return jsonify({'error': (
                f"\"Papkaoch\" bo'limida sozlangan manzil topilmadi: {folder_path}. "
                f"Shu bo'lim uchun alohida manzil kiriting yoki Papkaoch'dagi sozlamani tekshiring."
            )}), 400
        return jsonify({'error': f"Papka topilmadi yoki serverdan ko'rinmayapti: {folder_path}"}), 400

    # Sanani papka nomidan aniqlaymiz — yo'l ichidagi ENG YAQIN ota-papkalardan
    # birortasi "DD.MM.YYYY" ko'rinishida bo'lsa (masalan "05.08.2026"), o'shani
    # olamiz. MUHIM: faqat BEVOSITA ota-papkani emas, bir necha daraja yuqoridagi
    # ota-papkalarni ham tekshiramiz — chunki ba'zan sana papkasi bilan maqsad
    # papka ("Жалоба" va h.k.) orasida qo'shimcha oraliq papka(lar) bo'lishi mumkin
    # (masalan ".../05.08.2026/Biror papka nomi/Жалоба"). Eng yaqin (barglarga
    # eng yaqin) mos keluvchi papka ustunlik qiladi. Yoki hodim qo'lda kiritgan
    # qiymatdan aniqlanadi.
    sana = None
    if sana_override:
        try:
            datetime.strptime(sana_override, '%Y-%m-%d')
            sana = sana_override
        except ValueError:
            return jsonify({'error': "Sana formati noto'g'ri (YYYY-MM-DD kerak)"}), 400
    else:
        sana = _detect_sana_from_path(folder_path)

    folder_basename = os.path.basename(os.path.normpath(folder_path))
    name_warning = None
    if board == 't_hisobot':
        expected_names = COMMENT_FOLDER_NAMES
    else:
        expected_names = BOARD_SPECIAL_FOLDER_NAMES.get(board, ())
    if expected_names and folder_basename not in expected_names:
        name_warning = (
            f"Diqqat: papka nomi \"{folder_basename}\" — kutilgan nom(lar): "
            f"{', '.join(expected_names)}. Agar bu xato bo'lsa, manzilni tekshiring."
        )

    # Keyingi safar tezroq boshlash uchun shu manzilni eslab qolamiz
    # (hodim faqat sana qismini o'zgartirsa yetarli bo'ladi)
    set_setting(f'last_folder_path_{board}', folder_path)

    if board == 't_hisobot':
        sections = _build_t_hisobot_sections(folder_path, sana)
    else:
        sections = _build_jaloba_fishing_section(board, folder_path, sana)

    return jsonify({
        'ok': True,
        'board': board,
        'sana': sana,
        'sana_detected': bool(sana and not sana_override),
        'folder_path': folder_path,
        'name_warning': name_warning,
        'sections': sections,
    })

# --- Papkadan hisoblash: 2) TASDIQLASH (faqat shu bosqichda bazaga yoziladi) ---
# Jaloba/Fishing uchun qator: {link_id, order_num, url, count, sana, board}
# T.hisobot uchun qator:      {link_id, order_num, url, count (ajratilgan — papkadan hisoblangan), count2 (umumiy — qo'lda/mavjud), status, sana, board}
# link_id bo'lsa -> mavjud yozuv yangilanadi (faqat farq bo'lsa)
# link_id bo'lmasa -> hodim kiritgan URL bilan YANGI yozuv qo'shiladi
@app.route('/api/screenshot-folder/apply', methods=['POST'])
def apply_screenshot_folder():
    data = request.json or {}
    board = data.get('board', '')
    sana = data.get('sana', '')
    updates = data.get('updates', [])

    if board not in ('jaloba', 'fishing', 't_hisobot'):
        return jsonify({'error': "Noto'g'ri loyiha"}), 400
    try:
        datetime.strptime(sana, '%Y-%m-%d')
    except (ValueError, TypeError):
        return jsonify({'error': "Sana noto'g'ri yoki berilmagan"}), 400
    if not isinstance(updates, list) or not updates:
        return jsonify({'error': "Yangilanadigan ma'lumotlar bo'sh"}), 400

    added = 0
    updated = 0
    unchanged = 0
    errors = []
    is_th = board == 't_hisobot'

    with get_db() as conn:
        for u in updates:
            try:
                link_id_raw = u.get('link_id')
                order_num = int(u.get('order_num'))
                count = int(u.get('count') or 0)
                count2 = int(u.get('count2') or 0) if is_th else 0
                url = (u.get('url') or '').strip()
                row_status = valid_status_for_board(board, u.get('status', '')) if is_th else DEFAULT_STATUS.get(board, 'faol')
            except (TypeError, ValueError):
                errors.append(f"Noto'g'ri qiymat: {u}")
                continue

            if not url:
                errors.append(f"#{order_num}: link (URL) kiritilmagan — o'tkazib yuborildi")
                continue

            link_id = None
            if link_id_raw not in (None, '', 0, '0'):
                try:
                    link_id = int(link_id_raw)
                except (TypeError, ValueError):
                    link_id = None

            row = None
            if link_id:
                row = conn.execute('SELECT * FROM links WHERE id=?', (link_id,)).fetchone()
            if row is None:
                # Ehtiyot chorasi: agar shu board+url bilan boshqa yozuv allaqachon
                # mavjud bo'lsa (masalan tartib raqami boshqacha bo'lgan bo'lsa ham),
                # dublikat yaratmasdan o'shani yangilaymiz.
                row = conn.execute('SELECT * FROM links WHERE board=? AND url=?', (board, url)).fetchone()

            if row:
                is_new_day = (row['sana'] != sana)
                if is_th:
                    changed = (
                        row['url'] != url or
                        int(row['ajratilgan_izohlar'] or 0) != count or
                        int(row['umumiy_izohlar'] or 0) != count2 or
                        row['order_num'] != order_num or
                        row['sana'] != sana or
                        row['status'] != row_status
                    )
                    if changed:
                        edited_at = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                        if is_new_day:
                            # Boshqa kunga o'tyapti — o'sha linkka oldin (bugungidan
                            # oldin) yig'ilgan sonlarni prev_* ustunlariga qo'shib,
                            # yo'qotib yubormaymiz; joriy ustunlarga esa YANGI
                            # kunning sonini yozamiz.
                            new_prev_umumiy = int(row['prev_umumiy_izohlar'] or 0) + int(row['umumiy_izohlar'] or 0)
                            new_prev_ajratilgan = int(row['prev_ajratilgan_izohlar'] or 0) + int(row['ajratilgan_izohlar'] or 0)
                        else:
                            new_prev_umumiy = int(row['prev_umumiy_izohlar'] or 0)
                            new_prev_ajratilgan = int(row['prev_ajratilgan_izohlar'] or 0)
                        conn.execute(
                            'UPDATE links SET url=?, sana=?, ajratilgan_izohlar=?, umumiy_izohlar=?, '
                            'status=?, order_num=?, edited_at=?, prev_umumiy_izohlar=?, prev_ajratilgan_izohlar=? WHERE id=?',
                            (url, sana, count, count2, row_status, order_num, edited_at,
                             new_prev_umumiy, new_prev_ajratilgan, row['id'])
                        )
                        updated += 1
                    else:
                        unchanged += 1
                else:
                    changed = (
                        row['url'] != url or
                        int(row['likes'] or 0) != count or
                        row['order_num'] != order_num or
                        row['sana'] != sana
                    )
                    if changed:
                        edited_at = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                        if is_new_day:
                            new_prev_likes = int(row['prev_likes'] or 0) + int(row['likes'] or 0)
                        else:
                            new_prev_likes = int(row['prev_likes'] or 0)
                        conn.execute(
                            'UPDATE links SET url=?, sana=?, likes=?, order_num=?, edited_at=?, prev_likes=? WHERE id=?',
                            (url, sana, count, order_num, edited_at, new_prev_likes, row['id'])
                        )
                        updated += 1
                    else:
                        unchanged += 1
            else:
                # MUHIM: "created_at" bu yerda ANIQ (server lokal vaqti bilan)
                # yoziladi — SQLite'ning o'zidagi "DEFAULT CURRENT_TIMESTAMP"ga
                # ishonib qoldirilmaydi, chunki u UTC vaqtni yozadi, "edited_at"
                # esa quyida (va boshqa joylarda) doim datetime.now() — SERVER
                # LOKAL vaqti — bilan yoziladi. Ikkalasi turli vaqt mintaqasida
                # bo'lsa, "COALESCE(edited_at, created_at)" bo'yicha saralashda
                # (masalan T.hisobot bo'limida "eng oxirgi amal tepada" tartibi)
                # yangi yaratilgan qatorlar noto'g'ri joyga tushib qolishi mumkin.
                created_at = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                if is_th:
                    conn.execute(
                        'INSERT INTO links (url, sana, likes, status, board, order_num, ajratilgan_izohlar, umumiy_izohlar, created_at) '
                        'VALUES (?, ?, 0, ?, ?, ?, ?, ?, ?)',
                        (url, sana, row_status, board, order_num, count, count2, created_at)
                    )
                else:
                    conn.execute(
                        'INSERT INTO links (url, sana, likes, status, board, order_num, created_at) '
                        'VALUES (?, ?, ?, ?, ?, ?, ?)',
                        (url, sana, count, row_status, board, order_num, created_at)
                    )
                added += 1
        conn.commit()

    return jsonify({'ok': True, 'added': added, 'updated': updated, 'unchanged': unchanged, 'errors': errors})

@app.route('/api/links', methods=['GET'])
def get_links():
    clause, params = build_filters(request.args)

    # Standart bo'yicha: yangi kiritilgan yoki oxirgi tahrirlangan yozuv eng
    # yuqorida tursin (foydalanuvchi "eng oxirgi amal" ekanini darhol bilishi
    # uchun).
    sort_by = request.args.get('sort_by', 'last_activity')
    sort_dir = request.args.get('sort_dir', 'desc').lower()
    sort_col = SORT_COLUMNS.get(sort_by, 'last_activity')
    sort_dir_sql = 'ASC' if sort_dir == 'asc' else 'DESC'

    try:
        page = int(request.args.get('page', 1))
    except ValueError:
        page = 1
    try:
        per_page = int(request.args.get('per_page', 50))
    except ValueError:
        per_page = 50
    if per_page not in (50, 100, 500):
        per_page = 50
    if page < 1:
        page = 1

    with get_db() as conn:
        total = conn.execute(
            'SELECT COUNT(*) FROM links WHERE 1=1' + clause, params
        ).fetchone()[0]

        total_pages = max(1, math.ceil(total / per_page))
        if page > total_pages:
            page = total_pages
        offset = (page - 1) * per_page

        query = (
            'SELECT * FROM links WHERE 1=1' + clause +
            f' ORDER BY {sort_col} {sort_dir_sql}, id DESC LIMIT ? OFFSET ?'
        )
        rows = conn.execute(query, params + [per_page, offset]).fetchall()

    return jsonify({
        'items': [dict(r) for r in rows],
        'total': total,
        'page': page,
        'per_page': per_page,
        'total_pages': total_pages,
    })

def _group_rows_by_status(rows, board):
    """Qatorlar ro'yxatini status bo'yicha guruhlarga ajratadi, har bir guruh
    uchun jami sonlarni ham hisoblab qaytaradi. Bo'sh guruhlar (hech qanday
    qatori yo'q status) natijaga kiritilmaydi.

    T.HISOBOT UCHUN MUHIM: "Коммент" papkasi ichida oldindan belgilanmagan
    (ixtiyoriy nomdagi) rukun papkalari ham bo'lishi mumkin bo'lgani uchun,
    bu yerda faqat qat'iy STATUS_OPTIONS ro'yxati bilan CHEKLANMAYDI — aks
    holda yangi (masalan "Таълим" kabi) statuslar butunlay ko'rinmay
    qolardi. Buning o'rniga: avval taniqli statuslar (Umumiy, Tashrif, ...)
    odatdagi tartibda, keyin esa MA'LUMOTLARDA haqiqatan uchragan boshqa
    (yangi) statuslar birinchi uchragan tartibda qo'shiladi."""
    known_order = STATUS_OPTIONS.get(board, STATUS_OPTIONS[DEFAULT_BOARD])
    if board == 't_hisobot':
        seen = []
        for r in rows:
            if r['status'] not in seen:
                seen.append(r['status'])
        order = [s for s in known_order if s in seen] + [s for s in seen if s not in known_order]
    else:
        order = known_order

    groups = []
    for st in order:
        items = [r for r in rows if r['status'] == st]
        if not items:
            continue
        group = {'status': st, 'items': items, 'total': len(items)}
        if board == 't_hisobot':
            group['umumiy_izohlar_sum'] = sum(int(r['umumiy_izohlar'] or 0) for r in items)
            group['ajratilgan_izohlar_sum'] = sum(int(r['ajratilgan_izohlar'] or 0) for r in items)
        else:
            group['likes_sum'] = sum(int(r['likes'] or 0) for r in items)
        groups.append(group)
    return groups

@app.route('/api/links/grouped', methods=['GET'])
def get_links_grouped():
    """Joriy filtrlarga mos BARCHA yozuvlarni (sahifalashsiz) status bo'yicha
    guruhlab qaytaradi — sahifadagi va Wordga eksport qilingandagi kabi
    "har bir status uchun alohida jadval" ko'rinishini chizish uchun ishlatiladi."""
    board = get_board(request.args)
    clause, params = build_filters(request.args)

    sort_by = request.args.get('sort_by', 'last_activity')
    sort_dir = request.args.get('sort_dir', 'desc').lower()
    sort_col = SORT_COLUMNS.get(sort_by, 'last_activity')
    sort_dir_sql = 'ASC' if sort_dir == 'asc' else 'DESC'

    with get_db() as conn:
        query = (
            'SELECT * FROM links WHERE 1=1' + clause +
            f' ORDER BY {sort_col} {sort_dir_sql}, id DESC'
        )
        rows = [dict(r) for r in conn.execute(query, params).fetchall()]

    groups = _group_rows_by_status(rows, board)
    return jsonify({'groups': groups, 'total': len(rows)})

@app.route('/api/links', methods=['POST'])
def add_link():
    data = request.json
    board = data.get('board', DEFAULT_BOARD)
    board = board if board in BOARDS else DEFAULT_BOARD
    url = data.get('url', '').strip()
    sana = data.get('sana', '')
    likes = int(data.get('likes', 0) or 0)
    umumiy_izohlar = int(data.get('umumiy_izohlar', 0) or 0)
    ajratilgan_izohlar = int(data.get('ajratilgan_izohlar', 0) or 0)
    status = valid_status_for_board(board, data.get('status', ''))
    order_num_raw = data.get('order_num', None)
    if not url or not sana:
        return jsonify({'error': 'URL va sana majburiy'}), 400
    with get_db() as conn:
        # Bir xil link tekshiruvi — faqat shu loyiha (board) ichida
        existing = conn.execute(
            'SELECT * FROM links WHERE url=? AND board=?', (url, board)
        ).fetchone()
        if existing:
            return jsonify({'duplicate': True, 'existing': dict(existing)}), 409

        # Tartib raqami: agar berilmagan bo'lsa, shu board+sana (T.hisobot uchun —
        # board+sana+status, chunki u yerda raqamlash har bir status ichida
        # alohida boshlanadi) uchun avtomatik keyingi raqamni tanlaymiz
        if order_num_raw in (None, '', 0, '0'):
            if board == 't_hisobot':
                row = conn.execute(
                    'SELECT COALESCE(MAX(order_num), 0) as m FROM links WHERE board=? AND sana=? AND status=?',
                    (board, sana, status)
                ).fetchone()
            else:
                row = conn.execute(
                    'SELECT COALESCE(MAX(order_num), 0) as m FROM links WHERE board=? AND sana=?',
                    (board, sana)
                ).fetchone()
            order_num = (row['m'] or 0) + 1
        else:
            try:
                order_num = int(order_num_raw)
            except (TypeError, ValueError):
                order_num = None

        # MUHIM: created_at'ni ANIQ (server lokal vaqti bilan) yozamiz — pastdagi
        # "T.hisobot ustunda saralash" izohiga qarang (SQLite DEFAULT
        # CURRENT_TIMESTAMP — UTC, edited_at esa lokal vaqt — ikkalasi
        # aralashsa saralash buziladi).
        created_at = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        cur = conn.execute(
            'INSERT INTO links (url, sana, likes, status, board, umumiy_izohlar, ajratilgan_izohlar, order_num, created_at) '
            'VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)',
            (url, sana, likes, status, board, umumiy_izohlar, ajratilgan_izohlar, order_num, created_at)
        )
        conn.commit()
        row = conn.execute('SELECT * FROM links WHERE id=?', (cur.lastrowid,)).fetchone()
    return jsonify(dict(row)), 201

@app.route('/api/links/<int:lid>', methods=['PUT'])
def update_link(lid):
    data = request.json
    url = data.get('url', '').strip()
    sana = data.get('sana', '')
    likes = int(data.get('likes', 0) or 0)
    umumiy_izohlar = int(data.get('umumiy_izohlar', 0) or 0)
    ajratilgan_izohlar = int(data.get('ajratilgan_izohlar', 0) or 0)
    order_num_raw = data.get('order_num', None)
    with get_db() as conn:
        old = conn.execute('SELECT * FROM links WHERE id=?', (lid,)).fetchone()
        if not old:
            return jsonify({'error': 'Topilmadi'}), 404

        status = valid_status_for_board(old['board'], data.get('status', ''))

        if order_num_raw in (None, ''):
            order_num = old['order_num']
        else:
            try:
                order_num = int(order_num_raw)
            except (TypeError, ValueError):
                order_num = old['order_num']

        # Qatordagi biror maydon haqiqatan o'zgarganmi — shunga qarab
        # "o'zgartirilgan sana" (edited_at) bugungi sanaga yangilanadi,
        # aks holda avvalgi holicha (yoki bo'sh) qoladi
        changed = (
            old['url'] != url or
            old['sana'] != sana or
            int(old['likes']) != likes or
            int(old['umumiy_izohlar'] or 0) != umumiy_izohlar or
            int(old['ajratilgan_izohlar'] or 0) != ajratilgan_izohlar or
            old['status'] != status or
            (old['order_num'] or None) != (order_num or None)
        )
        edited_at = datetime.now().strftime('%Y-%m-%d %H:%M:%S') if changed else old['edited_at']

        conn.execute(
            'UPDATE links SET url=?, sana=?, likes=?, status=?, edited_at=?, '
            'umumiy_izohlar=?, ajratilgan_izohlar=?, order_num=? WHERE id=?',
            (url, sana, likes, status, edited_at, umumiy_izohlar, ajratilgan_izohlar, order_num, lid)
        )
        conn.commit()
        row = conn.execute('SELECT * FROM links WHERE id=?', (lid,)).fetchone()
    if not row:
        return jsonify({'error': 'Topilmadi'}), 404
    return jsonify(dict(row))

def _remove_shot_file(rel_path):
    """Diskdagi eski screenshot faylini xavfsiz o'chiradi (mavjud bo'lsa)."""
    if not rel_path:
        return
    full = os.path.join(app.static_folder, rel_path)
    full = os.path.normpath(full)
    # Papkadan tashqariga chiqmasligini tekshiramiz
    if not full.startswith(os.path.normpath(UPLOAD_DIR)):
        return
    if os.path.exists(full):
        try:
            os.remove(full)
        except OSError:
            pass

@app.route('/api/links/<int:lid>/screenshot', methods=['POST'])
def upload_screenshot(lid):
    """Havolaga tegishli 'dastlabki' yoki 'keyingi' (o'chgan) screenshot yuklaydi."""
    kind = request.form.get('kind', '')
    if kind not in SHOT_COLUMNS:
        return jsonify({'error': "Noto'g'ri tur (kind 'before' yoki 'after' bo'lishi kerak)"}), 400
    if 'file' not in request.files:
        return jsonify({'error': 'Fayl topilmadi'}), 400
    file = request.files['file']
    if not file.filename or not allowed_shot_file(file.filename):
        return jsonify({'error': 'Faqat rasm fayli qabul qilinadi (png, jpg, jpeg, webp, gif)'}), 400

    col = SHOT_COLUMNS[kind]
    with get_db() as conn:
        row = conn.execute('SELECT * FROM links WHERE id=?', (lid,)).fetchone()
        if not row:
            return jsonify({'error': 'Topilmadi'}), 404

        _remove_shot_file(row[col])

        ext = file.filename.rsplit('.', 1)[1].lower()
        fname = f"{lid}_{kind}_{uuid.uuid4().hex}.{ext}"
        file.save(os.path.join(UPLOAD_DIR, fname))
        rel_path = f"{UPLOAD_SUBDIR}/{fname}".replace(os.sep, '/')

        edited_at = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        conn.execute(f'UPDATE links SET {col}=?, edited_at=? WHERE id=?', (rel_path, edited_at, lid))
        conn.commit()
        row = conn.execute('SELECT * FROM links WHERE id=?', (lid,)).fetchone()

    return jsonify(dict(row))

@app.route('/api/links/<int:lid>/screenshot', methods=['DELETE'])
def delete_screenshot(lid):
    """Havolaga tegishli 'dastlabki' yoki 'keyingi' screenshotni o'chiradi."""
    kind = request.args.get('kind', '')
    if kind not in SHOT_COLUMNS:
        return jsonify({'error': "Noto'g'ri tur (kind 'before' yoki 'after' bo'lishi kerak)"}), 400

    col = SHOT_COLUMNS[kind]
    with get_db() as conn:
        row = conn.execute('SELECT * FROM links WHERE id=?', (lid,)).fetchone()
        if not row:
            return jsonify({'error': 'Topilmadi'}), 404

        _remove_shot_file(row[col])
        edited_at = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        conn.execute(f'UPDATE links SET {col}=NULL, edited_at=? WHERE id=?', (edited_at, lid))
        conn.commit()
        row = conn.execute('SELECT * FROM links WHERE id=?', (lid,)).fetchone()

    return jsonify(dict(row))

@app.route('/api/links/<int:lid>', methods=['DELETE'])
@admin_required
def delete_link(lid):
    with get_db() as conn:
        row = conn.execute('SELECT screenshot_before, screenshot_after FROM links WHERE id=?', (lid,)).fetchone()
        if row:
            _remove_shot_file(row['screenshot_before'])
            _remove_shot_file(row['screenshot_after'])
        conn.execute('DELETE FROM links WHERE id=?', (lid,))
        conn.commit()
    return jsonify({'ok': True})

@app.route('/api/links/bulk-delete', methods=['POST'])
@admin_required
def bulk_delete_links():
    """Belgilangan bir nechta havolani bittada o'chiradi. Faqat admin uchun."""
    data = request.json or {}
    ids = data.get('ids', [])
    if not isinstance(ids, list) or not ids:
        return jsonify({'error': "id'lar ro'yxati bo'sh"}), 400
    try:
        ids = [int(i) for i in ids]
    except (ValueError, TypeError):
        return jsonify({'error': "Noto'g'ri id qiymati"}), 400

    placeholders = ','.join('?' * len(ids))
    with get_db() as conn:
        rows = conn.execute(
            f'SELECT screenshot_before, screenshot_after FROM links WHERE id IN ({placeholders})', ids
        ).fetchall()
        for r in rows:
            _remove_shot_file(r['screenshot_before'])
            _remove_shot_file(r['screenshot_after'])
        cur = conn.execute(f'DELETE FROM links WHERE id IN ({placeholders})', ids)
        conn.commit()
        deleted = cur.rowcount
    return jsonify({'ok': True, 'deleted': deleted})

@app.route('/api/links/clear', methods=['POST'])
@admin_required
def clear_board_links():
    """Joriy loyiha (board) uchun jadvaldagi BARCHA yozuvlarni o'chiradi. Faqat admin uchun."""
    data = request.json or {}
    board = data.get('board', DEFAULT_BOARD)
    board = board if board in BOARDS else DEFAULT_BOARD
    with get_db() as conn:
        rows = conn.execute('SELECT screenshot_before, screenshot_after FROM links WHERE board=?', (board,)).fetchall()
        for r in rows:
            _remove_shot_file(r['screenshot_before'])
            _remove_shot_file(r['screenshot_after'])
        cur = conn.execute('DELETE FROM links WHERE board=?', (board,))
        conn.commit()
        deleted = cur.rowcount
    return jsonify({'ok': True, 'deleted': deleted, 'board': board})

@app.route('/api/stats', methods=['GET'])
def stats():
    """Umumiy statistika. Filtr parametrlari berilsa, natija shu filtrga mos hisoblanadi.
    T.hisobot boshqa ustunlarga ega bo'lgani uchun natija shakli board'ga qarab farqlanadi."""
    board = get_board(request.args)
    clause, params = build_filters(request.args)

    with get_db() as conn:
        total = conn.execute(
            'SELECT COUNT(*) FROM links WHERE 1=1' + clause, params
        ).fetchone()[0]

        if board == 't_hisobot':
            umumiy_sum = conn.execute(
                'SELECT COALESCE(SUM(umumiy_izohlar), 0) FROM links WHERE 1=1' + clause, params
            ).fetchone()[0]
            ajratilgan_sum = conn.execute(
                'SELECT COALESCE(SUM(ajratilgan_izohlar), 0) FROM links WHERE 1=1' + clause, params
            ).fetchone()[0]
            return jsonify({
                'mode': 't_hisobot',
                'total': total,
                'umumiy_izohlar_sum': umumiy_sum,
                'ajratilgan_izohlar_sum': ajratilgan_sum,
            })

        ochgan_status = "o'chgan"
        faol = conn.execute(
            'SELECT COUNT(*) FROM links WHERE status=?' + clause, ['faol'] + params
        ).fetchone()[0]
        ochgan = conn.execute(
            'SELECT COUNT(*) FROM links WHERE status=?' + clause, [ochgan_status] + params
        ).fetchone()[0]
        likes_sum = conn.execute(
            'SELECT COALESCE(SUM(likes), 0) FROM links WHERE 1=1' + clause, params
        ).fetchone()[0]

    return jsonify({'mode': 'jaloba', 'total': total, 'faol': faol, 'ochgan': ochgan, 'likes_sum': likes_sum})

# T.hisobot statuslari uchun ranglar (Excel eksportida qator foni)
T_HISOBOT_STATUS_COLORS = {
    'Umumiy':        'e2e8f0',
    'Tashrif':       'cfe8ff',
    'Videoselektor': 'e6d9ff',
    'Taqdimot':      'd2f4e6',
    'Tabrik':        'ffd9e8',
    'Saylov':        'ffe4bf',
}

def prev_total_text(current, prev):
    """Qavs ichida oldingi kunlargi jami, tashqarisida joriy (eng so'nggi kun)
    soni — masalan "(50) 10". Oldingi kunlar bo'lmasa — faqat son. Word/Excel
    eksportida shu matn ko'rinishida chiqadi; umumiy statistikaga esa faqat
    "joriy" (tashqaridagi) son qo'shiladi, qavs ichidagisi hisoblanmaydi."""
    c = int(current or 0)
    p = int(prev or 0)
    if p > 0:
        return f"({p}) {c}"
    return str(c)

@app.route('/api/export', methods=['GET'])
def export_excel():
    clause, params = build_filters(request.args)
    query = 'SELECT * FROM links WHERE 1=1' + clause + ' ORDER BY id'
    board = get_board(request.args)

    with get_db() as conn:
        rows = conn.execute(query, params).fetchall()

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Linklar'

    header_fill = openpyxl.styles.PatternFill('solid', fgColor='1a73e8')
    header_font = openpyxl.styles.Font(bold=True, color='FFFFFF')

    if board == 't_hisobot':
        headers = ['#', 'Link', 'Sana', 'Umumiy izohlar soni', 'Yozilgan izohlar soni', 'Status']
    else:
        headers = ['#', 'Link', 'Sana', 'Likelar', 'Status']

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = openpyxl.styles.Alignment(horizontal='center')

    green = openpyxl.styles.PatternFill('solid', fgColor='d4edda')
    red   = openpyxl.styles.PatternFill('solid', fgColor='f8d7da')

    for i, row in enumerate(rows, 1):
        ws.cell(row=i+1, column=1, value=i)
        ws.cell(row=i+1, column=2, value=row['url'])
        ws.cell(row=i+1, column=3, value=format_sana(row))

        if board == 't_hisobot':
            ws.cell(row=i+1, column=4, value=prev_total_text(row['umumiy_izohlar'], row['prev_umumiy_izohlar']))
            ws.cell(row=i+1, column=5, value=prev_total_text(row['ajratilgan_izohlar'], row['prev_ajratilgan_izohlar']))
            ws.cell(row=i+1, column=6, value=row['status'])
            fill_color = T_HISOBOT_STATUS_COLORS.get(row['status'], 'ffffff')
            fill = openpyxl.styles.PatternFill('solid', fgColor=fill_color)
            for col in range(1, 7):
                ws.cell(row=i+1, column=col).fill = fill
        else:
            ws.cell(row=i+1, column=4, value=prev_total_text(row['likes'], row['prev_likes']))
            ws.cell(row=i+1, column=5, value=row['status'])
            fill = green if row['status'] == 'faol' else red
            for col in range(1, 6):
                ws.cell(row=i+1, column=col).fill = fill

    ws.column_dimensions['A'].width = 6
    ws.column_dimensions['B'].width = 55
    ws.column_dimensions['C'].width = 18
    ws.column_dimensions['D'].width = 16
    ws.column_dimensions['E'].width = 16
    if board == 't_hisobot':
        ws.column_dimensions['F'].width = 16

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    filename = f"linklar_{board}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    return send_file(buf, as_attachment=True, download_name=filename,
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

def _set_cell_width(cell, width):
    cell.width = width
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = tcPr.makeelement(qn('w:tcW'), {})
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width.twips))
    tcW.set(qn('w:type'), 'dxa')

@app.route('/api/export/docx', methods=['GET'])
def export_docx():
    clause, params = build_filters(request.args)
    # Eng oxirgi kiritilgan yoki tahrirlangan yozuv jadvalning ENG OXIRIDA
    # tursin — shu sabab eng ESKI faoliyat vaqtidan boshlab (o'sish tartibida)
    # saralaymiz (veb sahifadagi "eng yangisi tepada" tartibiga qarama-qarshi).
    query = 'SELECT * FROM links WHERE 1=1' + clause + ' ORDER BY COALESCE(edited_at, created_at) ASC, id ASC'

    with get_db() as conn:
        rows = conn.execute(query, params).fetchall()

    board = get_board(request.args)
    is_thisobot = board == 't_hisobot'

    doc = Document()

    # Sahifani kitobiy (vertikal, portret) qilib sozlaymiz
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    heading = doc.add_heading(f"Havolalar ro'yxati — {BOARDS.get(board, board)}", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f"Yaratilgan sana: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.font.size = Pt(9)

    # Raqamli ustunlarning umumiy yig'indisi — jadval tepasida ko'rsatiladi
    if is_thisobot:
        total_umumiy = sum(int(r['umumiy_izohlar'] or 0) for r in rows)
        total_ajratilgan = sum(int(r['ajratilgan_izohlar'] or 0) for r in rows)
        totals_text = (
            f"Jami postlar: {len(rows)}  |  "
            f"Umumiy izohlar soni (jami): {total_umumiy}  |  "
            f"Yozilgan izohlar soni (jami): {total_ajratilgan}"
        )
    else:
        total_likes = sum(int(r['likes'] or 0) for r in rows)
        totals_text = f"Jami postlar: {len(rows)}  |  Jalobalar soni (jami): {total_likes}"

    totals_p = doc.add_paragraph()
    totals_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    totals_run = totals_p.add_run(totals_text)
    totals_run.bold = True
    totals_run.font.size = Pt(10)

    def add_screenshot_cell(cell, img_rel):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_full = os.path.join(app.static_folder, img_rel) if img_rel else None
        if img_full and os.path.exists(img_full):
            try:
                run = p.add_run()
                run.add_picture(img_full, width=Inches(1.0))
                return
            except Exception:
                pass
        p.add_run('—').font.size = Pt(9)

    def add_table_header(table, headers, col_widths):
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = ''
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            _set_cell_width(hdr_cells[i], col_widths[i])

    def new_table(headers, col_widths):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        add_table_header(table, headers, col_widths)
        return table

    def add_thisobot_row(table, col_widths, idx, row):
        cells = table.add_row().cells
        for i, w in enumerate(col_widths):
            _set_cell_width(cells[i], w)

        p = cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(str(idx)).font.size = Pt(9)

        p = cells[1].paragraphs[0]
        p.add_run(row['url'] or '').font.size = Pt(9)

        add_screenshot_cell(cells[2], row['screenshot_before'])

        p = cells[3].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(format_sana(row)).font.size = Pt(9)

        p = cells[4].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(prev_total_text(row['umumiy_izohlar'], row['prev_umumiy_izohlar'])).font.size = Pt(9)

        p = cells[5].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(prev_total_text(row['ajratilgan_izohlar'], row['prev_ajratilgan_izohlar'])).font.size = Pt(9)

        p = cells[6].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        status_run = p.add_run(row['status'] or '')
        status_run.font.size = Pt(9)
        status_run.bold = True

    def add_generic_row(table, col_widths, idx, row):
        cells = table.add_row().cells
        for i, w in enumerate(col_widths):
            _set_cell_width(cells[i], w)

        p = cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(str(idx)).font.size = Pt(9)

        add_screenshot_cell(cells[1], row['screenshot_before'])
        add_screenshot_cell(cells[2], row['screenshot_after'])

        p = cells[3].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(format_sana(row)).font.size = Pt(9)

        p = cells[4].paragraphs[0]
        p.add_run(row['url'] or '').font.size = Pt(9)

        p = cells[5].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(prev_total_text(row['likes'], row['prev_likes'])).font.size = Pt(9)

        p = cells[6].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        status_text = 'Faol' if row['status'] == 'faol' else "O'chgan"
        status_run = p.add_run(status_text)
        status_run.font.size = Pt(9)
        status_run.bold = True

    if is_thisobot:
        # T.hisobot uchun: har bir Status alohida jadvalga ajratiladi (STATUS_OPTIONS
        # tartibida), tepada esa har bir status bo'yicha alohida qisqa statistika
        # ("shundan, ... bo'yicha:") beriladi.
        headers = ['№', 'Link', 'Dastlabki', 'Sana', 'Umumiy izohlar soni', 'Yozilgan izohlar soni', 'Status']
        col_widths = [Inches(0.35), Inches(1.9), Inches(1.2), Inches(0.9), Inches(0.85), Inches(0.85), Inches(0.85)]

        groups = _group_rows_by_status(rows, board)

        if len(groups) > 1:
            shundan_p = doc.add_paragraph()
            shundan_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shundan_run = shundan_p.add_run('shundan,')
            shundan_run.bold = True
            shundan_run.font.size = Pt(10)

            for g in groups:
                gp_title = doc.add_paragraph()
                gp_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                gp_title_run = gp_title.add_run(f"{g['status']} bo'yicha:")
                gp_title_run.bold = True
                gp_title_run.font.size = Pt(10)

                gp_stats = doc.add_paragraph()
                gp_stats.alignment = WD_ALIGN_PARAGRAPH.CENTER
                gp_stats_run = gp_stats.add_run(
                    f"Jami postlar: {g['total']}  |  "
                    f"Umumiy izohlar soni (jami): {g['umumiy_izohlar_sum']}  |  "
                    f"Yozilgan izohlar soni (jami): {g['ajratilgan_izohlar_sum']}"
                )
                gp_stats_run.font.size = Pt(10)

        doc.add_paragraph()  # bo'shliq

        for g in groups:
            heading_p = doc.add_paragraph()
            heading_run = heading_p.add_run(g['status'])
            heading_run.bold = True
            heading_run.font.size = Pt(13)
            heading_run.font.color.rgb = RGBColor(0x2b, 0x2b, 0x2b)

            table = new_table(headers, col_widths)
            for idx, row in enumerate(g['items'], 1):
                add_thisobot_row(table, col_widths, idx, row)

            doc.add_paragraph()  # keyingi status jadvalidan oldin bo'shliq
    else:
        # Jaloba/Fishing — avvalgidek yagona jadval
        headers = ['№', 'Dastlabki screen', 'Keyingi screen', 'Sana', 'Link', "Jalobalar soni", 'Status']
        col_widths = [Inches(0.35), Inches(1.1), Inches(1.1), Inches(0.9), Inches(1.9), Inches(0.75), Inches(0.8)]

        doc.add_paragraph()  # jadval bilan orasida biroz bo'shliq
        table = new_table(headers, col_widths)
        for idx, row in enumerate(rows, 1):
            add_generic_row(table, col_widths, idx, row)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"linklar_{board}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return send_file(
        buf, as_attachment=True, download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/api/import', methods=['POST'])
def import_excel():
    if 'file' not in request.files:
        return jsonify({'error': 'Fayl topilmadi'}), 400
    file = request.files['file']
    if not file.filename.endswith(('.xlsx', '.xls')):
        return jsonify({'error': 'Faqat .xlsx yoki .xls fayl qabul qilinadi'}), 400

    board = request.form.get('board', DEFAULT_BOARD)
    board = board if board in BOARDS else DEFAULT_BOARD

    try:
        wb = openpyxl.load_workbook(BytesIO(file.read()), data_only=True)
        ws = wb.active
    except Exception as e:
        return jsonify({'error': f'Faylni o\'qib bo\'lmadi: {str(e)}'}), 400

    added = 0
    updated = 0
    unchanged = 0
    skipped = 0
    errors = []

    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return jsonify({'error': 'Fayl bo\'sh'}), 400

    header = [str(c).strip().lower() if c else '' for c in rows[0]]

    def find_col(names):
        for name in names:
            if name in header:
                return header.index(name)
        return None

    url_col  = find_col(['link', 'url', 'havola'])
    sana_col = find_col(['sana', 'date', 'kun'])
    like_col = find_col(['likelar', 'likes', 'like', 'jalobalar soni'])
    stat_col = find_col(['status', 'holat'])
    umumiy_col = find_col(['umumiy izohlar soni', 'umumiy izohlar', 'umumiy'])
    ajratilgan_col = find_col(['ajratilgan izohlar soni', 'ajratilgan izohlar', 'ajratilgan'])
    order_col = find_col(['tartib raqami', 'tartib', 'raqam', '№', 'no', 'order', 'order_num'])

    if url_col is None:
        url_col  = 1
        sana_col = 2
        like_col = 3
        stat_col = 4
        data_rows = rows
    else:
        data_rows = rows[1:]

    def parse_date(val):
        if val is None:
            return None
        if hasattr(val, 'strftime'):
            return val.strftime('%Y-%m-%d')
        s = str(val).strip()
        for fmt in ('%Y-%m-%d', '%d.%m.%Y', '%d/%m/%Y', '%m/%d/%Y', '%d-%m-%Y'):
            try:
                return datetime.strptime(s, fmt).strftime('%Y-%m-%d')
            except:
                pass
        return None

    def parse_count(val):
        try:
            return int(float(str(val))) if val not in (None, '') else 0
        except:
            return 0

    def parse_status(val, board):
        if val is None:
            options = STATUS_OPTIONS.get(board, STATUS_OPTIONS[DEFAULT_BOARD])
            return DEFAULT_STATUS.get(board, options[0])
        return valid_status_for_board(board, str(val))

    with get_db() as conn:
        # Board+sana bo'yicha "keyingi tartib raqami" keshi — har bir sanadagi
        # eng katta order_num'ni bazadan bir marta o'qib, shu import davomida
        # qo'shiladigan yangi qatorlarga ketma-ket raqam beramiz (har kuni 1 dan
        # boshlanadi, chunki kesh sana bo'yicha alohida saqlanadi).
        next_order_cache = {}

        def get_next_order(sana_key):
            if sana_key not in next_order_cache:
                row = conn.execute(
                    'SELECT COALESCE(MAX(order_num), 0) as m FROM links WHERE board=? AND sana=?',
                    (board, sana_key)
                ).fetchone()
                next_order_cache[sana_key] = (row['m'] or 0)
            next_order_cache[sana_key] += 1
            return next_order_cache[sana_key]

        for i, row in enumerate(data_rows, 2):
            try:
                url = str(row[url_col]).strip() if row[url_col] else ''
                if not url or url.lower() in ('none', 'url', 'link', 'havola', ''):
                    skipped += 1
                    continue

                sana = parse_date(row[sana_col]) if sana_col is not None and sana_col < len(row) else None
                if not sana:
                    sana = datetime.now().strftime('%Y-%m-%d')

                likes_raw = row[like_col] if like_col is not None and like_col < len(row) else 0
                likes = parse_count(likes_raw)

                umumiy_izohlar = parse_count(row[umumiy_col]) if umumiy_col is not None and umumiy_col < len(row) else 0
                ajratilgan_izohlar = parse_count(row[ajratilgan_col]) if ajratilgan_col is not None and ajratilgan_col < len(row) else 0

                status_raw = row[stat_col] if stat_col is not None and stat_col < len(row) else None
                status = parse_status(status_raw, board)

                order_raw = row[order_col] if order_col is not None and order_col < len(row) else None
                order_num_from_file = None
                if order_raw not in (None, ''):
                    try:
                        order_num_from_file = int(float(str(order_raw)))
                    except (ValueError, TypeError):
                        order_num_from_file = None

                # Bir xil havola (shu board ichida) allaqachon mavjudmi — tekshiramiz.
                # Mavjud bo'lsa: faqat haqiqatan o'zgargan maydonlarni yangilaymiz
                # (dublikat qator sifatida qayta qo'shmaymiz). Hech narsa o'zgarmagan
                # bo'lsa — umuman tegmaymiz.
                existing = conn.execute(
                    'SELECT * FROM links WHERE url=? AND board=?', (url, board)
                ).fetchone()

                if existing:
                    order_num = order_num_from_file if order_num_from_file is not None else existing['order_num']
                    changed = (
                        existing['sana'] != sana or
                        int(existing['likes'] or 0) != likes or
                        int(existing['umumiy_izohlar'] or 0) != umumiy_izohlar or
                        int(existing['ajratilgan_izohlar'] or 0) != ajratilgan_izohlar or
                        existing['status'] != status or
                        (existing['order_num'] or None) != (order_num or None)
                    )
                    if changed:
                        edited_at = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                        conn.execute(
                            'UPDATE links SET sana=?, likes=?, status=?, umumiy_izohlar=?, '
                            'ajratilgan_izohlar=?, order_num=?, edited_at=? WHERE id=?',
                            (sana, likes, status, umumiy_izohlar, ajratilgan_izohlar, order_num, edited_at, existing['id'])
                        )
                        updated += 1
                    else:
                        unchanged += 1
                    continue

                order_num = order_num_from_file if order_num_from_file is not None else get_next_order(sana)

                # MUHIM: created_at'ni ANIQ (server lokal vaqti bilan) yozamiz —
                # yuqoridagi izohlarga qarang (UTC/lokal vaqt aralashib ketishi
                # saralashni buzadi).
                created_at = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                conn.execute(
                    'INSERT INTO links (url, sana, likes, status, board, umumiy_izohlar, ajratilgan_izohlar, order_num, created_at) '
                    'VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)',
                    (url, sana, likes, status, board, umumiy_izohlar, ajratilgan_izohlar, order_num, created_at)
                )
                added += 1
            except Exception as e:
                errors.append(f'{i}-qator: {str(e)}')
                skipped += 1
        conn.commit()

    return jsonify({
        'added': added,
        'updated': updated,
        'unchanged': unchanged,
        'skipped': skipped,
        'errors': errors[:10],
    })


if __name__ == '__main__':
    init_db()
    start_papkaoch_bot()  # faqat TELEGRAM_BOT_TOKEN va TELEGRAM_GROUP_ID .env'da sozlangan bo'lsa ishga tushadi
    app.run(host='0.0.0.0', port=5003, debug=False, threaded=True)
